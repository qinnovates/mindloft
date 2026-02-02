#!/usr/bin/env python3
"""
ONI Framework: Organic Firewall Deep Dive
Expanding the security architecture for neural interfaces
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_firewall_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========== TITLE ==========
    title = doc.add_heading('', 0)
    title_run = title.add_run('ONI Framework: The Organic Firewall Architecture')
    title_run.font.size = Pt(22)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A Zero-Trust Security Model for Neural Interfaces')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = doc.add_paragraph()
    author.add_run('Kevin L. Qi').italic = True
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph("""This document expands on the Organic Network Interface (ONI) Framework's security architecture, specifically detailing the "Organic Firewall" concept. As brain-computer interfaces (BCIs) become bidirectional systems capable of both reading neural signals and writing stimulation patterns back to the brain, the attack surface expands dramatically.

We propose a multi-layered firewall architecture that treats BCI electrode arrays as edge nodes in a Zero-Trust model—where no signal (biological or digital) is trusted by default. This document details the physical implementation, logical architecture, and operational policies for securing the bio-digital boundary.""")

    # ========== SECTION 1 ==========
    doc.add_heading('1. The Problem: BCIs as Critical Attack Surfaces', level=1)

    doc.add_paragraph("""Modern BCIs like Neuralink's N1 implant are bidirectional systems:

• READ PATH (Egress): Neural signals → Amplification → Digitization → Compression → Wireless Transmission
• WRITE PATH (Ingress): External Commands → Wireless Reception → Validation → Electrical Stimulation → Neural Tissue

This bidirectionality creates unprecedented attack vectors. Recent research has identified critical threats:""")

    # Threat table
    doc.add_heading('1.1 Documented Attack Vectors', level=2)

    threat_table = doc.add_table(rows=1, cols=4)
    threat_table.style = 'Table Grid'

    headers = ['Attack Type', 'Vector', 'Impact', 'Range']
    for i, h in enumerate(headers):
        threat_table.rows[0].cells[i].text = h
        threat_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(threat_table.rows[0].cells[i], 'C00000')
        threat_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    threats = [
        ('Bluebugging', 'Bluetooth exploitation', 'Full device takeover, signal interception', '~10 meters'),
        ('Bluesnarfing', 'Unsecured BT connection', 'Neural data exfiltration', '~100 meters'),
        ('BlueBorne', 'BT stack vulnerability', 'Complete device control, malicious injection', 'Wireless range'),
        ('Thought Eavesdropping', 'Signal interception', 'Private memory/credential extraction', 'Network-wide'),
        ('Command Injection', 'Malicious stimulation', 'Involuntary movement, speech, actions', 'Direct'),
        ('Emotional Manipulation', 'Targeted stimulation', 'Fear, anxiety, pleasure center activation', 'Direct'),
        ('Neural Ransomware', 'Device lockout', 'Implant disabled until payment', 'Remote'),
        ('Man-in-the-Middle', 'Communication interception', 'Signal modification in transit', 'Wireless range'),
    ]

    for t in threats:
        row = threat_table.add_row()
        for i, val in enumerate(t):
            row.cells[i].text = val

    doc.add_paragraph()

    doc.add_heading('1.2 Why Traditional Security Models Fail', level=2)
    doc.add_paragraph("""Traditional perimeter security assumes a trusted internal network. This model fails catastrophically for BCIs because:

1. The "internal network" is living neural tissue — it cannot be patched, updated, or replaced
2. Compromise of the WRITE path causes immediate physical harm
3. The attack surface is literally inside the skull — physical isolation is impossible
4. Latency requirements (real-time neural processing) conflict with deep packet inspection
5. Power constraints (~25mW for N1) limit computational security overhead

This necessitates a Zero-Trust approach: verify every signal, in both directions, at every layer.""")

    # ========== SECTION 2 ==========
    doc.add_heading('2. BCI Nodes as Edge Nodes: The Network Topology', level=1)

    doc.add_paragraph("""In traditional networking, edge nodes are the ingress/egress points where traffic enters or exits a network boundary. In the ONI Framework, BCI electrode arrays serve exactly this function at the bio-digital boundary.""")

    doc.add_heading('2.1 ONI Network Topology', level=2)

    doc.add_paragraph("""The ONI architecture defines three network zones:""")

    zone_table = doc.add_table(rows=1, cols=4)
    zone_table.style = 'Table Grid'

    zone_headers = ['Zone', 'Components', 'Trust Level', 'ONI Layers']
    for i, h in enumerate(zone_headers):
        zone_table.rows[0].cells[i].text = h
        zone_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(zone_table.rows[0].cells[i], '2F5496')
        zone_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    zones = [
        ('ORGANIC ZONE', 'Neural tissue, neurons, synapses, neurotransmitters', 'PROTECTED (cannot be compromised without harm)', 'L9–L14'),
        ('EDGE ZONE', 'Electrode array, on-chip ASIC, signal processing', 'ZERO TRUST (primary inspection point)', 'L8'),
        ('DIGITAL ZONE', 'External pod, Bluetooth link, cloud services, applications', 'UNTRUSTED (assume compromised)', 'L1–L7'),
    ]

    for z in zones:
        row = zone_table.add_row()
        for i, val in enumerate(z):
            row.cells[i].text = val
            if i == 2:
                if 'PROTECTED' in val:
                    set_cell_shading(row.cells[i], 'E2F0D9')
                elif 'ZERO TRUST' in val:
                    set_cell_shading(row.cells[i], 'FFF2CC')
                elif 'UNTRUSTED' in val:
                    set_cell_shading(row.cells[i], 'F8CBAD')

    doc.add_paragraph()

    doc.add_heading('2.2 Traffic Flow Model', level=2)
    doc.add_paragraph("""
┌─────────────────────────────────────────────────────────────────────────┐
│                         ORGANIC ZONE (L9-L14)                           │
│   ┌─────────────────────────────────────────────────────────────────┐   │
│   │  Neural Tissue: Neurons, Synapses, Neurotransmitters            │   │
│   │  [Cognitive processes, memory, intent, identity]                │   │
│   └─────────────────────────────────────────────────────────────────┘   │
│                              ↑↓ (Action Potentials)                      │
└─────────────────────────────────────────────────────────────────────────┘
                               ↑↓
┌─────────────────────────────────────────────────────────────────────────┐
│                    EDGE ZONE - L8 (Neural Gateway)                      │
│  ┌────────────────────────────────────────────────────────────────────┐ │
│  │                    ╔══════════════════════════╗                    │ │
│  │                    ║   ORGANIC FIREWALL       ║                    │ │
│  │                    ║   ==================     ║                    │ │
│  │                    ║   • Signal Validation    ║                    │ │
│  │                    ║   • Anomaly Detection    ║                    │ │
│  │                    ║   • Rate Limiting        ║                    │ │
│  │                    ║   • Pattern Matching     ║                    │ │
│  │                    ║   • Encryption/Auth      ║                    │ │
│  │                    ║   • Audit Logging        ║                    │ │
│  │                    ╚══════════════════════════╝                    │ │
│  │                                                                    │ │
│  │  [Electrode Array] ←→ [ASIC Signal Processing] ←→ [Wireless TX]   │ │
│  │     (1024 ch)            (Amplify/Digitize)         (Bluetooth)   │ │
│  └────────────────────────────────────────────────────────────────────┘ │
└─────────────────────────────────────────────────────────────────────────┘
                               ↑↓ (Encrypted Bluetooth)
┌─────────────────────────────────────────────────────────────────────────┐
│                         DIGITAL ZONE (L1-L7)                            │
│   ┌─────────────┐    ┌─────────────┐    ┌─────────────────────────┐    │
│   │ External Pod│ ←→ │ Mobile App  │ ←→ │ Cloud Services / APIs   │    │
│   │ (Behind Ear)│    │ (Phone/PC)  │    │ (Updates, Analytics)    │    │
│   └─────────────┘    └─────────────┘    └─────────────────────────────┘    │
└─────────────────────────────────────────────────────────────────────────┘
""")

    # ========== SECTION 3 ==========
    doc.add_heading('3. Organic Firewall: Physical Architecture', level=1)

    doc.add_heading('3.1 Where Does the Firewall Physically Exist?', level=2)
    doc.add_paragraph("""The Organic Firewall is not a single component but a distributed security architecture spanning multiple physical locations:""")

    location_table = doc.add_table(rows=1, cols=4)
    location_table.style = 'Table Grid'

    loc_headers = ['Location', 'Physical Form', 'Function', 'Constraints']
    for i, h in enumerate(loc_headers):
        location_table.rows[0].cells[i].text = h
        location_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(location_table.rows[0].cells[i], '2F5496')
        location_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    locations = [
        ('PRIMARY: On-Implant', 'Dedicated silicon on N1 ASIC (~1mm²)', 'Real-time signal validation, rate limiting, anomaly detection', 'Power: <5mW, Latency: <1ms'),
        ('SECONDARY: External Pod', 'Processor in behind-ear device', 'Deep pattern analysis, encryption, authentication', 'Power: ~100mW, Latency: <10ms'),
        ('TERTIARY: Edge Gateway', 'Secure enclave on phone/PC', 'Policy enforcement, logging, cloud communication', 'Power: Unlimited, Latency: <100ms'),
        ('QUATERNARY: Cloud', 'Server-side security services', 'Threat intelligence, model updates, forensics', 'Power: N/A, Latency: Seconds'),
    ]

    for loc in locations:
        row = location_table.add_row()
        for i, val in enumerate(loc):
            row.cells[i].text = val

    doc.add_paragraph()

    doc.add_heading('3.2 On-Implant Firewall (Primary)', level=2)
    doc.add_paragraph("""The most critical firewall component resides on the implant itself. Given the Neuralink N1's specifications:

• Chip size: 4 × 5 mm
• Power budget: 24.7 mW total
• Channels: 1,024 electrodes
• Sample rate: 20,000 samples/second
• Data rate: ~200 Mbps raw, ~1 Mbps compressed

The on-implant firewall must operate within approximately 5mW and <1ms latency. This enables:""")

    p = doc.add_paragraph()
    p.add_run('EGRESS (READ) Path Protection:\n').bold = True
    p.add_run("""• Signal integrity validation (detect electrode degradation)
• Anomaly detection (unusual firing patterns)
• Data sanitization (prevent unintended information leakage)
• Compression with privacy preservation
• Encryption before wireless transmission
""")

    p = doc.add_paragraph()
    p.add_run('INGRESS (WRITE) Path Protection:\n').bold = True
    p.add_run("""• Command authentication (cryptographic verification)
• Stimulation pattern validation (safe amplitude/frequency bounds)
• Rate limiting (prevent overstimulation)
• Pattern matching against known attack signatures
• Emergency shutoff capability (hardware-enforced)
""")

    doc.add_heading('3.3 Physical Implementation Concept', level=2)
    doc.add_paragraph("""
┌────────────────────────────────────────────────────────────────────────────┐
│                     ON-IMPLANT ORGANIC FIREWALL                            │
│                        (Integrated into ASIC)                              │
├────────────────────────────────────────────────────────────────────────────┤
│                                                                            │
│  INGRESS PATH (Stimulation)              EGRESS PATH (Recording)           │
│  ─────────────────────────               ──────────────────────            │
│                                                                            │
│  [Bluetooth RX] ──┐                              ┌── [Bluetooth TX]        │
│                   ↓                              ↑                         │
│  ┌─────────────────────────┐      ┌─────────────────────────────┐         │
│  │ AUTHENTICATION MODULE   │      │  ENCRYPTION MODULE          │         │
│  │ • Verify command source │      │  • AES-256 encryption       │         │
│  │ • Check digital signature│      │  • Key rotation             │         │
│  │ • Session validation    │      │  • Nonce generation         │         │
│  └───────────┬─────────────┘      └──────────────┬──────────────┘         │
│              ↓                                   ↑                         │
│  ┌─────────────────────────┐      ┌─────────────────────────────┐         │
│  │ POLICY ENGINE           │      │  ANOMALY DETECTOR           │         │
│  │ • Allowed patterns DB   │      │  • Baseline neural patterns │         │
│  │ • Amplitude limits      │      │  • Statistical deviation    │         │
│  │ • Frequency bounds      │      │  • Seizure detection        │         │
│  │ • Rate limiting         │      │  • Privacy filter           │         │
│  └───────────┬─────────────┘      └──────────────┬──────────────┘         │
│              ↓                                   ↑                         │
│  ┌─────────────────────────┐      ┌─────────────────────────────┐         │
│  │ SAFETY VALIDATOR        │      │  SIGNAL CONDITIONER         │         │
│  │ • Hardware safety limits│      │  • Amplification            │         │
│  │ • Emergency shutoff     │      │  • Digitization (10-bit)    │         │
│  │ • Watchdog timer        │      │  • Filtering (500Hz-5kHz)   │         │
│  └───────────┬─────────────┘      └──────────────┬──────────────┘         │
│              ↓                                   ↑                         │
│         [STIMULATION                        [ELECTRODE                     │
│          CIRCUITS]                           ARRAY]                        │
│              ↓                                   ↑                         │
│  ════════════════════════════════════════════════════════════════         │
│                        NEURAL TISSUE (L9+)                                 │
│  ════════════════════════════════════════════════════════════════         │
│                                                                            │
│  ┌────────────────────────────────────────────────────────────┐           │
│  │                    AUDIT LOG (Secure Storage)              │           │
│  │  • All ingress commands (with timestamps)                  │           │
│  │  • Anomaly events                                          │           │
│  │  • Policy violations                                       │           │
│  │  • Tamper-resistant, cryptographically signed              │           │
│  └────────────────────────────────────────────────────────────┘           │
│                                                                            │
└────────────────────────────────────────────────────────────────────────────┘
""")

    # ========== SECTION 4 ==========
    doc.add_heading('4. Zero-Trust Implementation in ONI', level=1)

    doc.add_paragraph("""Traditional firewalls operate on "trust but verify" — allowing internal traffic while inspecting external traffic. Zero-Trust operates on "never trust, always verify" — treating ALL traffic as potentially hostile.""")

    doc.add_heading('4.1 Zero-Trust Principles Applied to ONI', level=2)

    zt_table = doc.add_table(rows=1, cols=3)
    zt_table.style = 'Table Grid'

    zt_headers = ['Zero-Trust Principle', 'Traditional Network', 'ONI Implementation']
    for i, h in enumerate(zt_headers):
        zt_table.rows[0].cells[i].text = h
        zt_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(zt_table.rows[0].cells[i], '2F5496')
        zt_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    zt_data = [
        ('Verify Explicitly', 'Authenticate users/devices', 'Authenticate ALL commands; verify neural signal authenticity'),
        ('Least Privilege Access', 'Role-based access control', 'Stimulation limited to specific regions; read access segmented'),
        ('Assume Breach', 'Monitor for lateral movement', 'Assume external pod compromised; implant validates independently'),
        ('Micro-segmentation', 'Network segmentation', 'Electrode groups isolated; cross-region stimulation requires escalation'),
        ('Continuous Validation', 'Session re-authentication', 'Every command re-verified; no persistent trust'),
        ('Encrypt Everything', 'TLS/SSL for traffic', 'All wireless communication encrypted; on-chip key storage'),
    ]

    for zt in zt_data:
        row = zt_table.add_row()
        for i, val in enumerate(zt):
            row.cells[i].text = val

    doc.add_paragraph()

    doc.add_heading('4.2 Defense-in-Depth Across ONI Layers', level=2)
    doc.add_paragraph("""Each layer of the ONI stack implements security controls:""")

    defense_table = doc.add_table(rows=1, cols=3)
    defense_table.style = 'Table Grid'

    def_headers = ['ONI Layer', 'Security Control', 'Implementation']
    for i, h in enumerate(def_headers):
        defense_table.rows[0].cells[i].text = h
        defense_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(defense_table.rows[0].cells[i], '2F5496')
        defense_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    defenses = [
        ('L1-L4 (Digital Transport)', 'Encryption, authentication', 'TLS 1.3, certificate pinning, secure boot'),
        ('L5-L7 (Digital Application)', 'Access control, input validation', 'OAuth 2.0, API rate limiting, anomaly detection'),
        ('L8 (Neural Gateway)', 'ORGANIC FIREWALL (Primary)', 'On-implant validation, hardware safety limits'),
        ('L9 (Ion Channel)', 'Stimulation bounds', 'Amplitude/frequency limits enforced in hardware'),
        ('L10 (Oscillatory)', 'Pattern validation', 'Prevent desynchronization attacks; rhythm enforcement'),
        ('L11-L12 (Cognitive)', 'Privacy preservation', 'On-device processing; differential privacy for transmitted data'),
        ('L13-L14 (Intent/Identity)', 'Agency protection', 'No direct WRITE access; read-only with consent'),
    ]

    for d in defenses:
        row = defense_table.add_row()
        for i, val in enumerate(d):
            row.cells[i].text = val
            if 'ORGANIC FIREWALL' in val:
                set_cell_shading(row.cells[i], 'FFF2CC')

    doc.add_paragraph()

    # ========== SECTION 5 ==========
    doc.add_heading('5. Firewall Policy Framework', level=1)

    doc.add_heading('5.1 Ingress (WRITE) Policies', level=2)
    doc.add_paragraph("""All stimulation commands must pass through the following policy chain:""")

    doc.add_paragraph("""
POLICY CHAIN: INGRESS (Stimulation Commands)
═══════════════════════════════════════════

[Incoming Command]
        ↓
┌───────────────────────────────────────┐
│ POLICY 1: AUTHENTICATION              │
│ ○ Valid cryptographic signature?      │
│ ○ Known authorized source?            │
│ ○ Session token valid?                │
│ FAIL → REJECT + LOG + ALERT           │
└───────────────────────────────────────┘
        ↓ PASS
┌───────────────────────────────────────┐
│ POLICY 2: AUTHORIZATION               │
│ ○ Source permitted for this region?   │
│ ○ Command type allowed for source?    │
│ ○ Time-of-day restrictions?           │
│ FAIL → REJECT + LOG                   │
└───────────────────────────────────────┘
        ↓ PASS
┌───────────────────────────────────────┐
│ POLICY 3: SAFETY BOUNDS               │
│ ○ Amplitude within safe range?        │
│ ○ Frequency within safe range?        │
│ ○ Duration within limits?             │
│ ○ Cumulative charge within limits?    │
│ FAIL → REJECT + LOG + ALERT           │
└───────────────────────────────────────┘
        ↓ PASS
┌───────────────────────────────────────┐
│ POLICY 4: RATE LIMITING               │
│ ○ Commands per second < threshold?    │
│ ○ Stimulation duty cycle < limit?     │
│ ○ Cool-down period respected?         │
│ FAIL → QUEUE or REJECT                │
└───────────────────────────────────────┘
        ↓ PASS
┌───────────────────────────────────────┐
│ POLICY 5: PATTERN MATCHING            │
│ ○ Pattern not in attack signature DB? │
│ ○ Pattern not anomalous vs baseline?  │
│ ○ Pattern consistent with intent?     │
│ FAIL → REJECT + LOG + ALERT           │
└───────────────────────────────────────┘
        ↓ PASS
┌───────────────────────────────────────┐
│ EXECUTE STIMULATION                   │
│ + Log command details                 │
│ + Monitor response                    │
│ + Update baseline                     │
└───────────────────────────────────────┘
""")

    doc.add_heading('5.2 Egress (READ) Policies', level=2)
    doc.add_paragraph("""
POLICY CHAIN: EGRESS (Neural Signal Transmission)
═════════════════════════════════════════════════

[Raw Neural Signals from Electrodes]
        ↓
┌───────────────────────────────────────┐
│ POLICY 1: SIGNAL VALIDATION           │
│ ○ Electrode impedance normal?         │
│ ○ Signal-to-noise ratio acceptable?   │
│ ○ No hardware fault indicators?       │
│ FAIL → FLAG channels + CONTINUE       │
└───────────────────────────────────────┘
        ↓
┌───────────────────────────────────────┐
│ POLICY 2: ANOMALY DETECTION           │
│ ○ Firing patterns within normal range?│
│ ○ No seizure-like activity?           │
│ ○ No external interference detected?  │
│ ANOMALY → ALERT + LOG + CONTINUE      │
└───────────────────────────────────────┘
        ↓
┌───────────────────────────────────────┐
│ POLICY 3: PRIVACY FILTERING           │
│ ○ Strip high-resolution raw data      │
│ ○ Apply differential privacy          │
│ ○ Aggregate before transmission       │
│ ○ Consent-based data categories       │
└───────────────────────────────────────┘
        ↓
┌───────────────────────────────────────┐
│ POLICY 4: COMPRESSION & ENCODING      │
│ ○ Compress (up to 200x)               │
│ ○ Format for transmission             │
│ ○ Add integrity checksum              │
└───────────────────────────────────────┘
        ↓
┌───────────────────────────────────────┐
│ POLICY 5: ENCRYPTION                  │
│ ○ Encrypt with session key            │
│ ○ Add authentication tag              │
│ ○ Sequence number (anti-replay)       │
└───────────────────────────────────────┘
        ↓
[Secure Bluetooth Transmission]
""")

    # ========== SECTION 6 ==========
    doc.add_heading('6. Emergency Protocols', level=1)

    doc.add_paragraph("""The Organic Firewall must include hardware-enforced emergency capabilities that cannot be overridden by software:""")

    emergency_table = doc.add_table(rows=1, cols=3)
    emergency_table.style = 'Table Grid'

    em_headers = ['Trigger Condition', 'Response', 'Recovery']
    for i, h in enumerate(em_headers):
        emergency_table.rows[0].cells[i].text = h
        emergency_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(emergency_table.rows[0].cells[i], 'C00000')
        emergency_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    emergencies = [
        ('Seizure-like activity detected', 'Immediate stimulation halt; alert clinician', 'Manual clinical review required'),
        ('Authentication failure threshold', 'Lock all WRITE access for 1 hour', 'Requires physical clinician presence'),
        ('Anomalous command pattern', 'Reject + isolate source + alert', 'Source must re-authenticate'),
        ('Hardware watchdog timeout', 'Safe-mode: READ-only operation', 'Requires device reset'),
        ('Cumulative charge limit exceeded', 'Stimulation disabled until reset', '24-hour cooldown or clinical override'),
        ('External magnet detected', 'Enter safe mode; disable wireless', 'Magnet removal + manual reset'),
    ]

    for e in emergencies:
        row = emergency_table.add_row()
        for i, val in enumerate(e):
            row.cells[i].text = val

    doc.add_paragraph()

    # ========== SECTION 7 ==========
    doc.add_heading('7. Regulatory Alignment', level=1)

    doc.add_paragraph("""The Organic Firewall architecture aligns with existing and emerging regulatory frameworks:""")

    reg_table = doc.add_table(rows=1, cols=3)
    reg_table.style = 'Table Grid'

    reg_headers = ['Regulation', 'Requirement', 'ONI Firewall Compliance']
    for i, h in enumerate(reg_headers):
        reg_table.rows[0].cells[i].text = h
        reg_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(reg_table.rows[0].cells[i], '2F5496')
        reg_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    regs = [
        ('FDA 21 CFR 820.30', 'Design controls for safety', 'Hardware safety limits; validation policies'),
        ('IEC 62443', 'Industrial cybersecurity', 'Defense-in-depth; zone segmentation'),
        ('HIPAA', 'Protected health information', 'Encryption; access controls; audit logs'),
        ('EU MDR', 'Medical device cybersecurity', 'Risk assessment; secure development'),
        ('Chile Neurodata Ruling', 'Neural data as human right', 'Privacy filtering; consent enforcement'),
        ('NIST CSF 2.0', 'Identify, Protect, Detect, Respond, Recover', 'Full framework implementation'),
    ]

    for r in regs:
        row = reg_table.add_row()
        for i, val in enumerate(r):
            row.cells[i].text = val

    doc.add_paragraph()

    # ========== SECTION 8 ==========
    doc.add_heading('8. Conclusion', level=1)

    doc.add_paragraph("""The Organic Firewall represents a fundamental requirement for safe BCI deployment. By treating electrode arrays as edge nodes in a Zero-Trust architecture, we can:

• Protect neural tissue from malicious stimulation
• Preserve privacy of neural data
• Maintain device integrity against sophisticated attacks
• Enable regulatory compliance
• Build public trust in neural interface technology

The firewall must be implemented primarily on-implant, with supporting layers in external devices. Hardware-enforced safety limits provide the ultimate backstop against software compromise.

As BCIs evolve from research tools to consumer medical devices, the Organic Firewall will become as essential as traditional network firewalls are today. The ONI Framework provides the architectural foundation for this critical security infrastructure.

"The brain's firewall is not optional—it is the minimum viable security for any system that touches living neural tissue."
""")

    # ========== REFERENCES ==========
    doc.add_heading('References', level=1)

    doc.add_paragraph("""1. Yale Digital Ethics Center. (2025). Study offers measures for safeguarding brain implants.

2. Black Cell Security. (2024). Threats of Thoughts: Cybersecurity Vulnerabilities of BCIs.

3. Schroder, T., et al. (2025). Cyber Risks to Next-Gen Brain-Computer Interfaces. arXiv.

4. Neuralink. (2021). 1024-Channel Simultaneous Recording Neural SoC.

5. NIST. (2024). Cybersecurity Framework 2.0.

6. FDA. (2021). Guidance for Brain-Computer Interface Devices.

7. Chile Supreme Court. (2024). Neurodata Privacy Ruling.

8. World Economic Forum. (2024). The BCI Market: Risks and Opportunities.""")

    # Save
    output_path = 'ONI Framework - Organic Firewall Deep Dive.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    create_firewall_document()
