#!/usr/bin/env python3
"""
Create the Organic Network Interface (ONI) Framework Research Paper
Updated naming convention: HAI/HMI → ONI
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_oni_document():
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========== TITLE ==========
    title = doc.add_heading('', 0)
    title_run = title.add_run('The Organic Network Interface (ONI) Framework')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A Unified Neuro-Computational Stack for Secure Bio-Digital Integration')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = doc.add_paragraph()
    author_run = author.add_run('Kevin L. Qi')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== ABSTRACT ==========
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph("""This paper proposes a unified layered framework—the Organic Network Interface (ONI) Model—for reasoning about the integration of biological neural systems with artificial intelligence. Extending the established OSI (Open Systems Interconnection) model beyond its traditional seven layers, we introduce seven additional organic layers (L8–L14) that characterize the brain's electrochemical, structural, and cognitive architecture.

The framework employs category-theoretic abstractions to formalize structure-preserving mappings between biological systems, AI architectures, and governance frameworks. By treating neural tissue as an organic network interface subject to entropy, attack surfaces, and coherence constraints, this work provides a defensible abstraction for identifying failure modes, security vulnerabilities, and governance checkpoints in emerging bio-digital systems.

This work does not claim to explain consciousness, nor does it attempt to unify neuroscience and artificial intelligence under a single mechanistic theory. Instead, it proposes a practical, structure-preserving framework compatible with existing compliance regimes such as FDA Class III medical device regulations. The ONI Framework is designed to be species-agnostic, applicable to any neural system—from animal research models to human clinical applications.""")

    # ========== SECTION 1 ==========
    doc.add_heading('1. Introduction', level=1)

    doc.add_heading('1.1 The Problem: No Framework for Organic Network Security', level=2)
    doc.add_paragraph("""As brain-computer interfaces (BCIs) transition from research prototypes to FDA-regulated medical devices, a critical gap emerges: while the OSI model provides a standardized framework for securing digital communications, no equivalent abstraction exists for the organic-digital boundary.

The trajectory of technology suggests an inevitable convergence. If we can already interface with neural tissue using wireless nodes (as demonstrated by Neuralink and similar technologies), and if all complex systems are subject to entropy and vulnerabilities, then we must anticipate attack vectors before they become exploitable zero-day vulnerabilities.

This paper addresses this gap by proposing the Organic Network Interface (ONI) Framework—a scale-aware, frequency-aware, and energy-aware communication stack that bridges machines and biological systems. The term "Organic Network Interface" was chosen deliberately to:

• Emphasize species-agnosticism: The framework applies equally to animal research models and human applications
• Parallel established IT terminology: "Network Interface" is immediately recognizable to the security and engineering communities
• Capture the biological substrate: "Organic" distinguishes living neural tissue from silicon-based systems""")

    doc.add_heading('1.2 Core Principles', level=2)
    doc.add_paragraph("""The ONI Framework is built on the following foundational principles:

1. Layered Abstraction: Each layer operates at characteristic frequencies, spatial scales, and energy profiles.

2. Scale Invariance: As we ascend the stack, frequency decreases while spatial scale and semantic compression increase.

3. Structure Preservation: Coherence—not mere signal transmission—is the fundamental invariant that must be maintained across layers.

4. Security by Design: Each layer boundary represents a potential attack surface requiring specific defensive measures.

5. Species Agnosticism: The framework applies to any neural system capable of interfacing with digital technology, from rodent models to primate studies to human clinical applications.""")

    # ========== SECTION 2 ==========
    doc.add_heading('2. Mathematical Foundations', level=1)

    doc.add_heading('2.1 The Scale-Frequency Relationship', level=2)
    doc.add_paragraph('A fundamental pattern emerges across all layers of the ONI stack: an inverse relationship between operational frequency and spatial scale. This can be expressed mathematically as:')

    formula1 = doc.add_paragraph()
    formula1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f1_run = formula1.add_run('f × S ≈ k')
    f1_run.font.size = Pt(14)
    f1_run.bold = True

    doc.add_paragraph("""Where:
• f = characteristic frequency (Hz)
• S = spatial scale (meters)
• k = scale-frequency constant (approximately 10⁶ m·Hz for neural systems)""")

    doc.add_heading('2.2 Coherence Metric', level=2)
    doc.add_paragraph('Signal coherence (Cₛ) represents the fundamental measure of information integrity across layers. We define coherence as:')

    formula2 = doc.add_paragraph()
    formula2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f2_run = formula2.add_run('Cₛ = e^(−(σ²ᵩ + σ²τ + σ²ᵧ))')
    f2_run.font.size = Pt(14)
    f2_run.bold = True

    doc.add_paragraph("""Where:
• σ²ᵩ = phase variance (timing jitter)
• σ²τ = transport variance (structural integrity)
• σ²ᵧ = gain variance (amplitude stability)

When Cₛ → 1, the signal maintains full structural integrity. When Cₛ → 0, information is effectively lost despite continued signal propagation.""")

    doc.add_heading('2.3 Category-Theoretic Formalization', level=2)
    doc.add_paragraph('Category theory provides the formal language for reasoning about structure-preserving mappings across the ONI framework. We define three primary categories:')

    doc.add_paragraph("""𝓑 (Biological Systems): Objects are neural structures; morphisms are signal propagation pathways.

𝓐 (AI/Artificial Systems): Objects are computational units; morphisms are information transformations.

𝓖 (Governance Systems): Objects are compliance checkpoints; morphisms are validation protocols.""")

    doc.add_paragraph("""
Functors and Natural Transformations:

• F: 𝓑 → 𝓐 maps biological coherence to AI architecture constraints
• G: 𝓐 → 𝓖 maps AI invariants to governance checkpoints
• η: F ⇒ G ensures coherence is preserved throughout the coevolution process""")

    # ========== SECTION 3 ==========
    doc.add_heading('3. The ONI Layered Model', level=1)
    doc.add_paragraph('The ONI Framework consists of 14 layers organized into three domains: the Digital Domain (L1–L7, traditional OSI), the Neural Interface Domain (L8–L10), and the Cognitive Domain (L11–L14). Each layer is characterized by its function, dominant signals, frequency range, spatial scale, and security considerations.')

    doc.add_heading('3.1 Design Principles for Layer Naming', level=2)
    doc.add_paragraph("""Layer names have been designed for intuitive comprehension while maintaining scientific accuracy. Each name reflects the layer's primary function using organic, accessible terminology:

• Lower layers use terms emphasizing physical and electrical properties
• Middle layers use terms emphasizing transport and encoding
• Higher layers use terms emphasizing cognition and identity

The naming convention follows the pattern: [Domain]-[Function] where applicable.""")

    doc.add_heading('3.2 Complete ONI Stack Reference', level=2)

    # Create comprehensive table
    layers = [
        ('L1', 'Physical Carrier', 'Electromagnetic transmission of raw bits', 'EM waves, voltage pulses, photons', 'Hz → THz', 'Nanometers → meters', 'Copper wire electrons, fiber optic photons, radio waves', 'digital'),
        ('L2', 'Link Framing', 'Error detection, frame synchronization', 'Symbols, frames, MAC addresses', 'MHz → GHz', 'Local network segments', 'Ethernet frames, WiFi packets, Bluetooth symbols', 'digital'),
        ('L3', 'Network Routing', 'Addressing, packet routing, path selection', 'IP packets, routing tables', 'kHz → MHz', 'LAN/WAN networks', 'IP addresses, BGP routes, OSPF paths', 'digital'),
        ('L4', 'Transport Flow', 'Reliable delivery, flow control, congestion', 'TCP/UDP segments, ports', 'Hz → kHz', 'End-to-end connections', 'TCP handshakes, UDP datagrams, port numbers', 'digital'),
        ('L5', 'Session State', 'Connection management, checkpointing', 'Session tokens, connection state', 'Seconds', 'Application endpoints', 'TLS sessions, RPC calls, NetBIOS names', 'digital'),
        ('L6', 'Data Encoding', 'Format translation, encryption, compression', 'Encoded bytes, encrypted payloads', 'Seconds', 'Application data units', 'SSL/TLS encryption, JPEG encoding, XML/JSON', 'digital'),
        ('L7', 'Application Interface', 'User-facing services, API endpoints', 'HTTP requests, application messages', 'Seconds → hours', 'Devices, servers, cloud', 'Web browsers, email clients, REST APIs', 'digital'),
        ('L8', 'Neural Gateway', 'Brain-machine physical boundary, electrode interface', 'Action potentials, local field potentials (LFPs)', '1–500 Hz', 'Neural tissue ↔ electrodes (μm–mm)', 'Utah arrays, Neuralink threads, ECoG grids, cochlear implants', 'neural'),
        ('L9', 'Ion Channel Encoding', 'Spike generation, neurotransmitter-digital translation', 'Ion channel dynamics (Na⁺, K⁺, Ca²⁺), spike trains', '10–200 Hz', 'Synaptic cleft (20–40 nm), microcircuits', 'Voltage-gated sodium channels, NMDA receptors, glutamate, GABA', 'neural'),
        ('L10', 'Oscillatory Synchronization', 'Temporal alignment, cross-region coordination', 'Brain oscillations (δ, θ, α, β, γ waves)', '0.5–100 Hz', 'Brain regions, cortical columns', 'Theta rhythms (hippocampus), Gamma bursts (attention), Alpha waves', 'neural'),
        ('L11', 'Cognitive Session', 'Working memory, attention windows, context maintenance', 'Persistent neural activation patterns', 'Seconds → minutes', 'Prefrontal cortex, parietal networks', 'Working memory traces, attention spotlight, task set representations', 'cognitive'),
        ('L12', 'Semantic Assembly', 'Concept formation, symbol binding, meaning construction', 'High-level distributed representations', 'Minutes → hours', 'Association cortex, temporal lobes', 'Word meanings, object concepts, semantic categories, mental imagery', 'cognitive'),
        ('L13', 'Intent & Agency', 'Goal formation, decision-making, value weighting', 'Executive cortex activity, reward signals', 'Hours → days', 'Prefrontal circuits, basal ganglia loops', 'Goal representations, action plans, reward prediction (dopamine)', 'cognitive'),
        ('L14', 'Identity & Ethics', 'Self-model, continuity of identity, moral reasoning', 'Integrated whole-brain patterns', 'Years → lifetime', 'Whole brain, autobiographical memory', 'Self-concept, personal values, ethical frameworks, life narrative', 'cognitive'),
    ]

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'

    headers = ['Layer', 'Name', 'Function', 'Dominant Signal', 'Frequency', 'Scale', 'Biological Examples']
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2F5496')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    domain_colors = {'digital': 'D6DCE5', 'neural': 'E2F0D9', 'cognitive': 'FFF2CC'}

    for layer in layers:
        row = table.add_row()
        for i in range(7):
            row.cells[i].text = layer[i]
            set_cell_shading(row.cells[i], domain_colors[layer[7]])

    doc.add_paragraph()
    legend = doc.add_paragraph()
    legend.add_run('Table Legend: ').bold = True
    legend.add_run('Blue-gray = Digital Domain (OSI L1-L7) | Green = Neural Interface Domain (L8-L10) | Yellow = Cognitive Domain (L11-L14)')

    # ========== SECTION 4 ==========
    doc.add_heading('4. Detailed Layer Analysis with Biological Examples', level=1)

    doc.add_heading('4.1 Neural Interface Domain (L8–L10)', level=2)

    p = doc.add_paragraph()
    p.add_run('Layer 8: Neural Gateway').bold = True

    doc.add_paragraph("""The Neural Gateway represents the critical boundary between silicon and carbon—the physical interface where digital signals meet biological tissue. This layer is analogous to the physical layer (L1) but operates at biological frequencies. In the ONI Framework, L8 serves as the primary ingress/egress point—analogous to how a fortress gate is both the entry point and the most critical point to defend.

Key Biological Components:
• Electrode arrays (Utah arrays: 96-channel, 400μm spacing)
• Neuralink N1 chip (1,024 electrodes per thread)
• Electrocorticography (ECoG) grids
• Peripheral nerve interfaces

Signal Characteristics:
• Action potentials: ~100 mV amplitude, 1-2 ms duration
• Local field potentials: 10-1000 μV, 1-500 Hz
• Signal-to-noise ratio: typically 5:1 to 20:1

Research Model Considerations:
• Rodent models: Smaller electrode arrays, higher spatial density requirements
• Primate models: Closer anatomical analogy to human applications
• Human applications: FDA Class III regulatory pathway""")

    p = doc.add_paragraph()
    p.add_run('Layer 9: Ion Channel Encoding').bold = True

    doc.add_paragraph("""This layer handles the translation between neural spike codes and interpretable digital representations. Ion channels serve as the biological equivalent of transistors—gating information flow based on voltage thresholds.

Key Biological Components:
• Voltage-gated Na⁺ channels (rapid depolarization, ~1 ms)
• Voltage-gated K⁺ channels (repolarization, ~2-4 ms)
• Ca²⁺ channels (neurotransmitter release trigger)
• NMDA receptors (coincidence detection, learning)
• AMPA receptors (fast excitatory transmission)

Neurotransmitter Examples at Different Timescales:
• Glutamate: Fast excitation (1-10 ms postsynaptic response)
• GABA: Fast inhibition (1-10 ms)
• Acetylcholine: Modulation (10-100 ms)
• Dopamine: Reward signaling (100 ms phasic, minutes tonic)
• Serotonin: Mood modulation (seconds to hours)""")

    p = doc.add_paragraph()
    p.add_run('Layer 10: Oscillatory Synchronization').bold = True

    doc.add_paragraph("""Neural oscillations provide the temporal coordination necessary for information routing across brain regions. Different frequency bands serve distinct cognitive functions.

Oscillation Frequency Bands:
• Delta (δ): 0.5–4 Hz — Deep sleep, cortical inhibition
• Theta (θ): 4–8 Hz — Memory encoding, spatial navigation
• Alpha (α): 8–12 Hz — Relaxed wakefulness, inhibitory gating
• Beta (β): 13–30 Hz — Active thinking, motor planning
• Gamma (γ): 30–100 Hz — Attention, sensory binding, consciousness

Cross-Frequency Coupling:
The theta-gamma code in the hippocampus exemplifies hierarchical temporal organization: 4-8 distinct gamma cycles (representing individual memory items) nest within each theta cycle, enabling sequential encoding of multiple items.""")

    doc.add_heading('4.2 Cognitive Domain (L11–L14)', level=2)

    doc.add_paragraph("""The Cognitive Domain encompasses the higher-order processing that transforms raw neural activity into thought, intention, and identity.

Layer 11 (Cognitive Session) Examples:
• Prefrontal cortex maintaining task rules
• Parietal cortex tracking spatial attention
• Working memory capacity: 4±1 items (Cowan's limit)

Layer 12 (Semantic Assembly) Examples:
• Temporal lobe semantic representations
• Wernicke's area language comprehension
• Concept cells (e.g., "Jennifer Aniston neurons")

Layer 13 (Intent & Agency) Examples:
• Dorsolateral prefrontal cortex: goal maintenance
• Anterior cingulate cortex: conflict monitoring
• Dopaminergic reward prediction (VTA → NAcc pathway)

Layer 14 (Identity & Ethics) Examples:
• Default Mode Network: self-referential processing
• Medial prefrontal cortex: self-concept
• Hippocampal autobiographical memory
• Orbitofrontal cortex: moral reasoning""")

    # ========== SECTION 5 ==========
    doc.add_heading('5. Security Framework: Organic Tactics, Techniques, and Procedures', level=1)

    doc.add_paragraph('Extending the MITRE ATT&CK framework concept to biological systems, we propose an Organic TTP taxonomy for anticipating and defending against attacks on neural interfaces. This framework applies across all research and clinical applications of the ONI model.')

    doc.add_heading('5.1 Attack Surface by Layer', level=2)

    # Security table
    sec_table = doc.add_table(rows=1, cols=4)
    sec_table.style = 'Table Grid'

    sec_headers = ['Layer', 'Attack Vector', 'Example Technique', 'Defense Strategy']
    sec_header_cells = sec_table.rows[0].cells
    for i, h in enumerate(sec_headers):
        sec_header_cells[i].text = h
        sec_header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(sec_header_cells[i], '2F5496')
        sec_header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    sec_data = [
        ('L8', 'Signal Injection', 'Malicious stimulation patterns', 'Hardware attestation, signal validation'),
        ('L9', 'Encoding Manipulation', 'Spike train replay attacks', 'Temporal signatures, anomaly detection'),
        ('L10', 'Desynchronization', 'Disrupting gamma rhythms', 'Phase-locking monitoring, rhythm enforcement'),
        ('L11', 'Cognitive Disruption', 'Attention hijacking', 'Working memory integrity checks'),
        ('L12', 'Semantic Inference', 'Thought reconstruction', 'On-device processing, differential privacy'),
        ('L13', 'Goal Manipulation', 'Reward signal hacking', 'Agency verification, intent logging'),
        ('L14', 'Identity Compromise', 'Personality alteration', 'Long-term pattern monitoring, ethical bounds'),
    ]

    for data in sec_data:
        row = sec_table.add_row()
        for i, val in enumerate(data):
            row.cells[i].text = val

    doc.add_paragraph()

    doc.add_heading('5.2 The Organic Firewall Concept', level=2)

    doc.add_paragraph("""Analogous to network firewalls that inspect and filter traffic at network boundaries, an "organic firewall" would operate at L8—the Neural Gateway—implementing:

1. Signal Validation: Verifying that incoming stimulation patterns match authorized signatures
2. Anomaly Detection: Identifying abnormal neural activity patterns that may indicate attack or malfunction
3. Access Control: Enforcing policies about which systems can read from or write to specific neural regions
4. Audit Logging: Maintaining tamper-resistant records of all interface interactions

This architecture follows Zero Trust principles: no signal is trusted by default, regardless of origin. The ONI Framework positions L8 as the critical control point—the "gate" of the organic fortress that requires the most robust defenses.""")

    # ========== SECTION 6 ==========
    doc.add_heading('6. Governance and Regulatory Alignment', level=1)

    doc.add_paragraph("""The ONI Framework is designed to be compatible with existing regulatory frameworks, particularly FDA Class III medical device requirements for brain-computer interfaces.

Key compliance mappings:

• FDA 21 CFR Part 820 (Quality System Regulation): Layer-specific quality controls
• IEC 62304 (Medical Device Software): Software lifecycle requirements for L1-L7
• ISO 14971 (Risk Management): Attack surface analysis per layer
• HIPAA: Data protection requirements for L11-L14 cognitive data
• IACUC Protocols: Animal research compliance for pre-clinical ONI studies

The category-theoretic formalization enables automated compliance verification through functorial mappings between system states and regulatory requirements. Importantly, the ONI Framework's species-agnostic design allows consistent governance from animal research through human clinical trials.""")

    # ========== SECTION 7 ==========
    doc.add_heading('7. Limitations and Non-Claims', level=1)

    doc.add_paragraph("""To avoid overextension and misinterpretation, we explicitly state the boundaries of this work:

1. Not a Theory of Consciousness: This framework does not attempt to explain subjective experience, qualia, or the ontological nature of consciousness.

2. Not a Biological Mechanism Claim: References to biological structures illustrate structural roles and constraints, not causality in cognitive processing.

3. Not Empirically Validated: The framework guides future empirical work and system design but does not replace controlled experimentation.

4. Not a Unifying Physical Theory: Category theory is employed as a formal language for structure-preserving mappings, not as an assertion about the fundamental nature of cognition.

5. No Guaranteed Safety: While the framework enables reasoning about failure modes, it does not guarantee alignment or safety under all conditions.

6. Species-Specific Variations: While the ONI Framework is designed to be species-agnostic, specific parameter values (frequencies, timescales, spatial scales) will vary across species and must be calibrated empirically.

In summary, this framework claims usefulness, not completeness; structure, not ontology; and risk reduction, not elimination.""")

    # ========== SECTION 8 ==========
    doc.add_heading('8. Future Work', level=1)

    doc.add_paragraph("""1. Organic TTP Scenario Modeling: Develop scenario-based attack simulations analogous to cyber kill chains.

2. Biological Signature Identification: Define measurable markers correlated with coherence loss or adversarial influence.

3. Detection and Response Policies: Create organic detection playbooks analogous to SOC procedures.

4. AI Safety Integration: Enable joint monitoring of biological and artificial agents for co-evolutionary governance.

5. Validation Through Animal Models: Test the framework using established neuroscience research models before human application.

6. Cross-Species Calibration: Establish ONI parameter mappings across commonly used research species (rodent, primate, human).""")

    # ========== SECTION 9 ==========
    doc.add_heading('9. Conclusion', level=1)

    doc.add_paragraph("""The Organic Network Interface (ONI) Framework provides a structured abstraction for reasoning about the integration of biological and artificial intelligence systems. By extending the OSI model with seven organic layers characterized by frequency, scale, and coherence properties, the framework enables:

• Systematic identification of attack surfaces across the bio-digital boundary
• Formal reasoning about structure preservation through category-theoretic mappings
• Alignment with existing regulatory and compliance frameworks
• Proactive defense against emerging threats to neural interfaces
• Consistent methodology from animal research through human clinical applications

As brain-computer interfaces transition from research to regulated medical devices, frameworks like ONI become essential for ensuring that bio-digital integration proceeds safely and ethically—regardless of the species under study.

"Intelligence—biological or artificial—fails not when signals disappear, but when structure fails. This is our framework for the future."
""")

    # ========== REFERENCES ==========
    doc.add_heading('References', level=1)

    doc.add_paragraph("""1. Buzsáki, G. (2006). Rhythms of the Brain. Oxford University Press.

2. Mac Lane, S. (1998). Categories for the Working Mathematician. Springer.

3. Dehaene, S., & Changeux, J. P. (2011). Experimental and theoretical approaches to conscious processing. Neuron, 70(2), 200-227.

4. Friston, K. (2010). The free-energy principle: a unified brain theory? Nature Reviews Neuroscience, 11(2), 127-138.

5. Marr, D. (1982). Vision: A Computational Investigation. MIT Press.

6. Musk, E., & Neuralink. (2019). An integrated brain-machine interface platform. Journal of Medical Internet Research.

7. MITRE Corporation. ATT&CK Framework. https://attack.mitre.org/

8. FDA. (2021). Guidance for Brain-Computer Interface Devices. 21 CFR Part 820.

9. Tononi, G., & Koch, C. (2015). Consciousness: here, there and everywhere? Philosophical Transactions of the Royal Society B.

10. ISO/IEC. (2023). Artificial Intelligence Risk Management Framework.""")

    # Save document
    output_path = 'The Organic Network Interface (ONI) Framework.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    create_oni_document()
