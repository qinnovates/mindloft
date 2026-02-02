#!/usr/bin/env python3
"""
ONI Framework: Organic Firewall Deep Dive
Version 2 - Using tables for diagrams instead of ASCII
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, border_color="000000", border_size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def merge_cells_in_row(table, row_idx, start_col, end_col):
    """Merge cells in a row from start_col to end_col"""
    row = table.rows[row_idx]
    start_cell = row.cells[start_col]
    for i in range(start_col + 1, end_col + 1):
        start_cell.merge(row.cells[i])
    return start_cell

def create_firewall_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========== TITLE ==========
    title = doc.add_heading('', 0)
    title_run = title.add_run('ONI Framework: The Organic Firewall Architecture')
    title_run.font.size = Pt(22)
    title_run.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('A Zero-Trust Security Model for Neural Interfaces')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = doc.add_paragraph()
    author.add_run('Kevin L. Qi').italic = True
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph("""This document expands on the Organic Network Interface (ONI) Framework's security architecture, specifically detailing the "Organic Firewall" concept. As brain-computer interfaces (BCIs) become bidirectional systems capable of both reading neural signals and writing stimulation patterns back to the brain, the attack surface expands dramatically.

We propose a multi-layered firewall architecture that treats BCI electrode arrays as edge nodes in a Zero-Trust model—where no signal (biological or digital) is trusted by default. This document details the physical implementation, logical architecture, and operational policies for securing the bio-digital boundary.""")

    # ========== SECTION 1 ==========
    doc.add_heading('1. The Problem: BCIs as Critical Attack Surfaces', level=1)

    doc.add_paragraph("""Modern BCIs like Neuralink's N1 implant are bidirectional systems:""")

    # Bidirectional paths table
    path_table = doc.add_table(rows=2, cols=2)
    path_table.style = 'Table Grid'

    path_table.rows[0].cells[0].text = "READ PATH (Egress)"
    path_table.rows[0].cells[1].text = "WRITE PATH (Ingress)"
    set_cell_shading(path_table.rows[0].cells[0], '2E75B6')
    set_cell_shading(path_table.rows[0].cells[1], 'C00000')
    path_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    path_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    path_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    path_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    path_table.rows[1].cells[0].text = "Neural Signals → Amplification → Digitization → Compression → Wireless TX"
    path_table.rows[1].cells[1].text = "External Commands → Wireless RX → Validation → Electrical Stimulation → Neural Tissue"
    set_cell_shading(path_table.rows[1].cells[0], 'DEEBF7')
    set_cell_shading(path_table.rows[1].cells[1], 'FBE5D6')

    doc.add_paragraph()
    doc.add_paragraph("This bidirectionality creates unprecedented attack vectors. Recent research has identified critical threats:")

    # Threat table
    doc.add_heading('1.1 Documented Attack Vectors', level=2)

    threat_table = doc.add_table(rows=1, cols=4)
    threat_table.style = 'Table Grid'

    headers = ['Attack Type', 'Vector', 'Impact', 'Range']
    for i, h in enumerate(headers):
        threat_table.rows[0].cells[i].text = h
        threat_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(threat_table.rows[0].cells[i], 'C00000')
        threat_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    threats = [
        ('Bluebugging', 'Bluetooth exploitation', 'Full device takeover, signal interception', '~10 meters'),
        ('Bluesnarfing', 'Unsecured BT connection', 'Neural data exfiltration', '~100 meters'),
        ('BlueBorne', 'BT stack vulnerability', 'Complete device control, malicious injection', 'Wireless range'),
        ('Thought Eavesdropping', 'Signal interception', 'Private memory/credential extraction', 'Network-wide'),
        ('Command Injection', 'Malicious stimulation', 'Involuntary movement, speech, actions', 'Direct'),
        ('Emotional Manipulation', 'Targeted stimulation', 'Fear, anxiety, pleasure center activation', 'Direct'),
        ('Neural Ransomware', 'Device lockout', 'Implant disabled until payment', 'Remote'),
        ('Man-in-the-Middle', 'Communication interception', 'Signal modification in transit', 'Wireless range'),
    ]

    for t in threats:
        row = threat_table.add_row()
        for i, val in enumerate(t):
            row.cells[i].text = val

    doc.add_paragraph()

    doc.add_heading('1.2 Why Traditional Security Models Fail', level=2)
    doc.add_paragraph("""Traditional perimeter security assumes a trusted internal network. This model fails catastrophically for BCIs because:

1. The "internal network" is living neural tissue — it cannot be patched, updated, or replaced
2. Compromise of the WRITE path causes immediate physical harm
3. The attack surface is literally inside the skull — physical isolation is impossible
4. Latency requirements (real-time neural processing) conflict with deep packet inspection
5. Power constraints (~25mW for N1) limit computational security overhead

This necessitates a Zero-Trust approach: verify every signal, in both directions, at every layer.""")

    # ========== SECTION 2 ==========
    doc.add_heading('2. BCI Nodes as Edge Nodes: The Network Topology', level=1)

    doc.add_paragraph("""In traditional networking, edge nodes are the ingress/egress points where traffic enters or exits a network boundary. In the ONI Framework, BCI electrode arrays serve exactly this function at the bio-digital boundary.""")

    doc.add_heading('2.1 ONI Network Zones', level=2)

    zone_table = doc.add_table(rows=4, cols=4)
    zone_table.style = 'Table Grid'

    zone_headers = ['Zone', 'Components', 'Trust Level', 'ONI Layers']
    for i, h in enumerate(zone_headers):
        zone_table.rows[0].cells[i].text = h
        zone_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(zone_table.rows[0].cells[i], '2F5496')
        zone_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    zones = [
        ('ORGANIC ZONE', 'Neural tissue, neurons, synapses, neurotransmitters', 'PROTECTED', 'L9–L14', 'E2F0D9'),
        ('EDGE ZONE', 'Electrode array, on-chip ASIC, signal processing', 'ZERO TRUST', 'L8', 'FFF2CC'),
        ('DIGITAL ZONE', 'External pod, Bluetooth, cloud services, apps', 'UNTRUSTED', 'L1–L7', 'F8CBAD'),
    ]

    for idx, z in enumerate(zones):
        row = zone_table.rows[idx + 1]
        row.cells[0].text = z[0]
        row.cells[1].text = z[1]
        row.cells[2].text = z[2]
        row.cells[3].text = z[3]
        for cell in row.cells:
            set_cell_shading(cell, z[4])
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ========== NETWORK TOPOLOGY DIAGRAM (Table-based) ==========
    doc.add_heading('2.2 Traffic Flow Architecture', level=2)

    # Create a visual diagram using nested tables
    topo_table = doc.add_table(rows=7, cols=1)
    topo_table.style = 'Table Grid'

    # Row 0: Organic Zone Header
    cell = topo_table.rows[0].cells[0]
    cell.text = "ORGANIC ZONE (L9-L14)"
    set_cell_shading(cell, '70AD47')
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 1: Organic Zone Content
    cell = topo_table.rows[1].cells[0]
    cell.text = "Neural Tissue: Neurons • Synapses • Neurotransmitters\nCognitive Processes: Memory • Intent • Identity • Agency"
    set_cell_shading(cell, 'E2F0D9')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 2: Arrow
    cell = topo_table.rows[2].cells[0]
    cell.text = "↑↓ Action Potentials (Bidirectional)"
    set_cell_shading(cell, 'FFFFFF')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True

    # Row 3: Edge Zone Header
    cell = topo_table.rows[3].cells[0]
    cell.text = "EDGE ZONE — L8: Neural Gateway — ORGANIC FIREWALL"
    set_cell_shading(cell, 'ED7D31')
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 4: Edge Zone Content
    cell = topo_table.rows[4].cells[0]
    set_cell_shading(cell, 'FFF2CC')

    # Create nested table for firewall components
    edge_content = """Electrode Array (1024 ch) → ASIC Signal Processing → Organic Firewall → Wireless TX/RX

FIREWALL FUNCTIONS:
• Signal Validation  • Anomaly Detection  • Rate Limiting
• Pattern Matching  • Encryption/Auth  • Audit Logging"""
    cell.text = edge_content
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 5: Arrow
    cell = topo_table.rows[5].cells[0]
    cell.text = "↑↓ Encrypted Bluetooth (Bidirectional)"
    set_cell_shading(cell, 'FFFFFF')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].bold = True

    # Row 6: Digital Zone
    cell = topo_table.rows[6].cells[0]
    cell.text = "DIGITAL ZONE (L1-L7)"
    set_cell_shading(cell, '5B9BD5')
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Digital zone detail table
    digital_table = doc.add_table(rows=1, cols=3)
    digital_table.style = 'Table Grid'

    digital_components = [
        ('External Pod\n(Behind Ear)', 'Mobile App\n(Phone/PC)', 'Cloud Services\n(Updates, Analytics)')
    ]

    for i, comp in enumerate(digital_components[0]):
        digital_table.rows[0].cells[i].text = comp
        set_cell_shading(digital_table.rows[0].cells[i], 'DEEBF7')
        digital_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== SECTION 3 ==========
    doc.add_heading('3. Organic Firewall: Physical Architecture', level=1)

    doc.add_heading('3.1 Firewall Location Hierarchy', level=2)
    doc.add_paragraph("The Organic Firewall is distributed across multiple physical locations:")

    location_table = doc.add_table(rows=5, cols=4)
    location_table.style = 'Table Grid'

    loc_headers = ['Location', 'Physical Form', 'Function', 'Constraints']
    for i, h in enumerate(loc_headers):
        location_table.rows[0].cells[i].text = h
        location_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(location_table.rows[0].cells[i], '2F5496')
        location_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    locations = [
        ('PRIMARY:\nOn-Implant', 'Dedicated silicon on N1 ASIC (~1mm²)', 'Real-time signal validation, rate limiting, anomaly detection', 'Power: <5mW\nLatency: <1ms', 'FFF2CC'),
        ('SECONDARY:\nExternal Pod', 'Processor in behind-ear device', 'Deep pattern analysis, encryption, authentication', 'Power: ~100mW\nLatency: <10ms', 'DEEBF7'),
        ('TERTIARY:\nEdge Gateway', 'Secure enclave on phone/PC', 'Policy enforcement, logging, cloud communication', 'Power: Unlimited\nLatency: <100ms', 'E2F0D9'),
        ('QUATERNARY:\nCloud', 'Server-side security services', 'Threat intelligence, model updates, forensics', 'Power: N/A\nLatency: Seconds', 'D9D9D9'),
    ]

    for idx, loc in enumerate(locations):
        row = location_table.rows[idx + 1]
        for i in range(4):
            row.cells[i].text = loc[i]
            set_cell_shading(row.cells[i], loc[4])
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ========== ON-IMPLANT FIREWALL DIAGRAM ==========
    doc.add_heading('3.2 On-Implant Firewall Architecture', level=2)

    doc.add_paragraph("""The most critical firewall component resides on the implant itself. Given the Neuralink N1's specifications (4×5mm chip, 24.7mW power, 1,024 channels), the on-implant firewall must operate within approximately 5mW and <1ms latency.""")

    # Create firewall architecture as table
    fw_table = doc.add_table(rows=3, cols=2)
    fw_table.style = 'Table Grid'

    # Headers
    fw_table.rows[0].cells[0].text = "INGRESS PATH (Stimulation Commands)"
    fw_table.rows[0].cells[1].text = "EGRESS PATH (Neural Signal Recording)"
    set_cell_shading(fw_table.rows[0].cells[0], 'C00000')
    set_cell_shading(fw_table.rows[0].cells[1], '2E75B6')
    fw_table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    fw_table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    fw_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    fw_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    fw_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    fw_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Flow content
    ingress_flow = """[Bluetooth RX]
    ↓
AUTHENTICATION MODULE
• Verify command source
• Check digital signature
• Session validation
    ↓
POLICY ENGINE
• Allowed patterns DB
• Amplitude limits
• Frequency bounds
• Rate limiting
    ↓
SAFETY VALIDATOR
• Hardware safety limits
• Emergency shutoff
• Watchdog timer
    ↓
[STIMULATION CIRCUITS]
    ↓
═══ NEURAL TISSUE ═══"""

    egress_flow = """[Bluetooth TX]
    ↑
ENCRYPTION MODULE
• AES-256 encryption
• Key rotation
• Nonce generation
    ↑
ANOMALY DETECTOR
• Baseline neural patterns
• Statistical deviation
• Seizure detection
• Privacy filter
    ↑
SIGNAL CONDITIONER
• Amplification
• Digitization (10-bit)
• Filtering (500Hz-5kHz)
    ↑
[ELECTRODE ARRAY]
    ↑
═══ NEURAL TISSUE ═══"""

    fw_table.rows[1].cells[0].text = ingress_flow
    fw_table.rows[1].cells[1].text = egress_flow
    set_cell_shading(fw_table.rows[1].cells[0], 'FBE5D6')
    set_cell_shading(fw_table.rows[1].cells[1], 'DEEBF7')

    # Audit log row
    audit_cell = fw_table.rows[2].cells[0]
    audit_cell.merge(fw_table.rows[2].cells[1])
    audit_cell.text = "AUDIT LOG (Secure Storage): All commands • Anomaly events • Policy violations • Tamper-resistant, cryptographically signed"
    set_cell_shading(audit_cell, 'D9D9D9')
    audit_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== SECTION 4 ==========
    doc.add_heading('4. Zero-Trust Implementation in ONI', level=1)

    doc.add_paragraph("""Traditional firewalls operate on "trust but verify" — allowing internal traffic while inspecting external traffic. Zero-Trust operates on "never trust, always verify" — treating ALL traffic as potentially hostile.""")

    doc.add_heading('4.1 Zero-Trust Principles Applied to ONI', level=2)

    zt_table = doc.add_table(rows=7, cols=3)
    zt_table.style = 'Table Grid'

    zt_headers = ['Zero-Trust Principle', 'Traditional Network', 'ONI Implementation']
    for i, h in enumerate(zt_headers):
        zt_table.rows[0].cells[i].text = h
        zt_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(zt_table.rows[0].cells[i], '2F5496')
        zt_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    zt_data = [
        ('Verify Explicitly', 'Authenticate users/devices', 'Authenticate ALL commands; verify neural signal authenticity'),
        ('Least Privilege Access', 'Role-based access control', 'Stimulation limited to specific regions; read access segmented'),
        ('Assume Breach', 'Monitor for lateral movement', 'Assume external pod compromised; implant validates independently'),
        ('Micro-segmentation', 'Network segmentation', 'Electrode groups isolated; cross-region stimulation requires escalation'),
        ('Continuous Validation', 'Session re-authentication', 'Every command re-verified; no persistent trust'),
        ('Encrypt Everything', 'TLS/SSL for traffic', 'All wireless communication encrypted; on-chip key storage'),
    ]

    for idx, zt in enumerate(zt_data):
        row = zt_table.rows[idx + 1]
        for i, val in enumerate(zt):
            row.cells[i].text = val

    doc.add_paragraph()

    doc.add_heading('4.2 Defense-in-Depth Across ONI Layers', level=2)

    defense_table = doc.add_table(rows=8, cols=3)
    defense_table.style = 'Table Grid'

    def_headers = ['ONI Layer', 'Security Control', 'Implementation']
    for i, h in enumerate(def_headers):
        defense_table.rows[0].cells[i].text = h
        defense_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(defense_table.rows[0].cells[i], '2F5496')
        defense_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    defenses = [
        ('L1-L4\n(Digital Transport)', 'Encryption, authentication', 'TLS 1.3, certificate pinning, secure boot', 'DEEBF7'),
        ('L5-L7\n(Digital Application)', 'Access control, input validation', 'OAuth 2.0, API rate limiting, anomaly detection', 'DEEBF7'),
        ('L8\n(Neural Gateway)', 'ORGANIC FIREWALL', 'On-implant validation, hardware safety limits', 'FFF2CC'),
        ('L9\n(Ion Channel)', 'Stimulation bounds', 'Amplitude/frequency limits enforced in hardware', 'E2F0D9'),
        ('L10\n(Oscillatory)', 'Pattern validation', 'Prevent desynchronization attacks; rhythm enforcement', 'E2F0D9'),
        ('L11-L12\n(Cognitive)', 'Privacy preservation', 'On-device processing; differential privacy', 'E2F0D9'),
        ('L13-L14\n(Intent/Identity)', 'Agency protection', 'No direct WRITE access; read-only with consent', 'E2F0D9'),
    ]

    for idx, d in enumerate(defenses):
        row = defense_table.rows[idx + 1]
        row.cells[0].text = d[0]
        row.cells[1].text = d[1]
        row.cells[2].text = d[2]
        for cell in row.cells:
            set_cell_shading(cell, d[3])
        if 'ORGANIC FIREWALL' in d[1]:
            row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ========== SECTION 5 ==========
    doc.add_heading('5. Firewall Policy Chains', level=1)

    doc.add_heading('5.1 Ingress (WRITE) Policy Chain', level=2)
    doc.add_paragraph("All stimulation commands must pass through the following policy chain:")

    # Ingress policy chain as table
    ingress_policy = doc.add_table(rows=6, cols=3)
    ingress_policy.style = 'Table Grid'

    policies_in = [
        ('POLICY 1', 'AUTHENTICATION', '• Valid cryptographic signature?\n• Known authorized source?\n• Session token valid?', 'REJECT + LOG + ALERT'),
        ('POLICY 2', 'AUTHORIZATION', '• Source permitted for this region?\n• Command type allowed?\n• Time-of-day restrictions?', 'REJECT + LOG'),
        ('POLICY 3', 'SAFETY BOUNDS', '• Amplitude within safe range?\n• Frequency within safe range?\n• Duration/charge within limits?', 'REJECT + LOG + ALERT'),
        ('POLICY 4', 'RATE LIMITING', '• Commands/second < threshold?\n• Stimulation duty cycle < limit?\n• Cool-down period respected?', 'QUEUE or REJECT'),
        ('POLICY 5', 'PATTERN MATCHING', '• Not in attack signature DB?\n• Not anomalous vs baseline?\n• Consistent with stated intent?', 'REJECT + LOG + ALERT'),
        ('✓ PASS', 'EXECUTE STIMULATION', '• Log command details\n• Monitor response\n• Update baseline', '—'),
    ]

    for idx, p in enumerate(policies_in):
        row = ingress_policy.rows[idx]
        row.cells[0].text = p[0]
        row.cells[1].text = p[1]
        row.cells[2].text = p[2]

        if idx < 5:
            set_cell_shading(row.cells[0], 'C00000')
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            set_cell_shading(row.cells[1], 'FBE5D6')
            set_cell_shading(row.cells[2], 'FBE5D6')
        else:
            set_cell_shading(row.cells[0], '70AD47')
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            set_cell_shading(row.cells[1], 'E2F0D9')
            set_cell_shading(row.cells[2], 'E2F0D9')

        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[1].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    doc.add_heading('5.2 Egress (READ) Policy Chain', level=2)

    egress_policy = doc.add_table(rows=6, cols=2)
    egress_policy.style = 'Table Grid'

    policies_out = [
        ('POLICY 1: SIGNAL VALIDATION', '• Electrode impedance normal?\n• Signal-to-noise acceptable?\n• No hardware fault indicators?'),
        ('POLICY 2: ANOMALY DETECTION', '• Firing patterns within normal range?\n• No seizure-like activity?\n• No external interference?'),
        ('POLICY 3: PRIVACY FILTERING', '• Strip high-resolution raw data\n• Apply differential privacy\n• Consent-based data categories'),
        ('POLICY 4: COMPRESSION', '• Compress (up to 200x)\n• Format for transmission\n• Add integrity checksum'),
        ('POLICY 5: ENCRYPTION', '• Encrypt with session key\n• Add authentication tag\n• Sequence number (anti-replay)'),
        ('→ SECURE TRANSMISSION', 'Encrypted Bluetooth to External Pod'),
    ]

    for idx, p in enumerate(policies_out):
        row = egress_policy.rows[idx]
        row.cells[0].text = p[0]
        row.cells[1].text = p[1]

        if idx < 5:
            set_cell_shading(row.cells[0], '2E75B6')
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            set_cell_shading(row.cells[1], 'DEEBF7')
        else:
            set_cell_shading(row.cells[0], '70AD47')
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
            set_cell_shading(row.cells[1], 'E2F0D9')

        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ========== SECTION 6 ==========
    doc.add_heading('6. Emergency Protocols', level=1)

    doc.add_paragraph("The Organic Firewall includes hardware-enforced emergency capabilities that cannot be overridden by software:")

    emergency_table = doc.add_table(rows=7, cols=3)
    emergency_table.style = 'Table Grid'

    em_headers = ['Trigger Condition', 'Automatic Response', 'Recovery Procedure']
    for i, h in enumerate(em_headers):
        emergency_table.rows[0].cells[i].text = h
        emergency_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(emergency_table.rows[0].cells[i], 'C00000')
        emergency_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    emergencies = [
        ('Seizure-like activity detected', 'Immediate stimulation halt; alert clinician', 'Manual clinical review required'),
        ('Authentication failure threshold (3x)', 'Lock all WRITE access for 1 hour', 'Requires physical clinician presence'),
        ('Anomalous command pattern', 'Reject + isolate source + alert', 'Source must re-authenticate'),
        ('Hardware watchdog timeout', 'Safe-mode: READ-only operation', 'Requires device reset'),
        ('Cumulative charge limit exceeded', 'Stimulation disabled until reset', '24-hour cooldown or clinical override'),
        ('External magnet detected', 'Enter safe mode; disable wireless', 'Magnet removal + manual reset'),
    ]

    for idx, e in enumerate(emergencies):
        row = emergency_table.rows[idx + 1]
        for i, val in enumerate(e):
            row.cells[i].text = val
            set_cell_shading(row.cells[i], 'FBE5D6')

    doc.add_paragraph()

    # ========== SECTION 7 ==========
    doc.add_heading('7. Regulatory Alignment', level=1)

    reg_table = doc.add_table(rows=7, cols=3)
    reg_table.style = 'Table Grid'

    reg_headers = ['Regulation', 'Requirement', 'ONI Firewall Compliance']
    for i, h in enumerate(reg_headers):
        reg_table.rows[0].cells[i].text = h
        reg_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(reg_table.rows[0].cells[i], '2F5496')
        reg_table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)

    regs = [
        ('FDA 21 CFR 820.30', 'Design controls for safety', 'Hardware safety limits; validation policies'),
        ('IEC 62443', 'Industrial cybersecurity', 'Defense-in-depth; zone segmentation'),
        ('HIPAA', 'Protected health information', 'Encryption; access controls; audit logs'),
        ('EU MDR', 'Medical device cybersecurity', 'Risk assessment; secure development'),
        ('Chile Neurodata Ruling', 'Neural data as human right', 'Privacy filtering; consent enforcement'),
        ('NIST CSF 2.0', 'Identify, Protect, Detect, Respond, Recover', 'Full framework implementation'),
    ]

    for idx, r in enumerate(regs):
        row = reg_table.rows[idx + 1]
        for i, val in enumerate(r):
            row.cells[i].text = val

    doc.add_paragraph()

    # ========== SECTION 8 ==========
    doc.add_heading('8. Conclusion', level=1)

    doc.add_paragraph("""The Organic Firewall represents a fundamental requirement for safe BCI deployment. By treating electrode arrays as edge nodes in a Zero-Trust architecture, we can:

• Protect neural tissue from malicious stimulation
• Preserve privacy of neural data
• Maintain device integrity against sophisticated attacks
• Enable regulatory compliance
• Build public trust in neural interface technology

The firewall must be implemented primarily on-implant, with supporting layers in external devices. Hardware-enforced safety limits provide the ultimate backstop against software compromise.

As BCIs evolve from research tools to consumer medical devices, the Organic Firewall will become as essential as traditional network firewalls are today. The ONI Framework provides the architectural foundation for this critical security infrastructure.""")

    conclusion_box = doc.add_table(rows=1, cols=1)
    conclusion_box.style = 'Table Grid'
    cell = conclusion_box.rows[0].cells[0]
    cell.text = '"The brain\'s firewall is not optional—it is the minimum viable security for any system that touches living neural tissue."'
    set_cell_shading(cell, '2F5496')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    cell.paragraphs[0].runs[0].italic = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ========== REFERENCES ==========
    doc.add_heading('References', level=1)

    doc.add_paragraph("""1. Yale Digital Ethics Center. (2025). Study offers measures for safeguarding brain implants.

2. Black Cell Security. (2024). Threats of Thoughts: Cybersecurity Vulnerabilities of BCIs.

3. Schroder, T., et al. (2025). Cyber Risks to Next-Gen Brain-Computer Interfaces. arXiv.

4. Neuralink. (2021). 1024-Channel Simultaneous Recording Neural SoC.

5. NIST. (2024). Cybersecurity Framework 2.0.

6. FDA. (2021). Guidance for Brain-Computer Interface Devices.

7. Chile Supreme Court. (2024). Neurodata Privacy Ruling.

8. World Economic Forum. (2024). The BCI Market: Risks and Opportunities.""")

    # Save
    output_path = 'ONI Framework - Organic Firewall Deep Dive.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')

if __name__ == '__main__':
    create_firewall_document()
