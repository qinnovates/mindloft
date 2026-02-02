#!/usr/bin/env python3
"""
Create Medium Blog Post: Week 2 - Neural Ransomware
Formatted for Medium - Accessible to general audience
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import textwrap

def clean_text(text):
    """Remove leading whitespace from each line"""
    return textwrap.dedent(text).strip()

def add_separator(doc):
    """Add a three-dot separator (Medium style)"""
    p = doc.add_paragraph()
    p.add_run('• • •')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_medium_post():
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========== TITLE ==========
    title = doc.add_heading('', 0)
    title_run = title.add_run('Neural Ransomware Isn\'t Science Fiction')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('The technical kill chain for holding a brain implant hostage — and why we need to talk about it now.')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    add_separator(doc)

    # ========== HOOK ==========
    hook = clean_text("""
    You wake up and something is wrong.

    The neural implant that restored your vision after the accident — the one that lets you see your children's faces — is displaying a message directly into your visual cortex:

    "Your device has been locked. Send 2 BTC to the following address within 48 hours or functionality will be permanently disabled."

    You can't see. You can't drive. You can't work. And somewhere, an attacker is waiting for payment.

    This isn't science fiction. Every component of this attack exists today.
    """)
    doc.add_paragraph(hook)

    add_separator(doc)

    # ========== SECTION: Not Hypothetical ==========
    doc.add_heading('This Attack Is Already Possible', level=1)

    p1 = clean_text("""
    Let me walk you through exactly how it would work.

    Neuralink's N1 chip — the one already implanted in human patients — communicates via Bluetooth Low Energy. The same protocol in your wireless headphones. The same protocol that's been exploited by attacks like BlueBorne, which compromised billions of devices in 2017.

    BlueBorne didn't require pairing. It didn't require user interaction. If your Bluetooth was on, you were vulnerable.

    Now imagine that vulnerability in something you can't turn off. Something inside your skull.
    """)
    doc.add_paragraph(p1)

    add_separator(doc)

    # ========== SECTION: The Kill Chain ==========
    doc.add_heading('The Kill Chain: From Bluetooth to Brainjack', level=1)

    p2 = clean_text("""
    Security researchers use "kill chains" to map out attack sequences. Here's what a neural ransomware attack might look like:
    """)
    doc.add_paragraph(p2)

    # Step 1
    s1 = doc.add_paragraph()
    s1.add_run('Step 1: Reconnaissance. ').bold = True
    s1.add_run('The attacker identifies targets. Maybe they\'re scanning for BCI-specific Bluetooth signatures in a hospital waiting room. Maybe they bought a list of patients from a compromised medical database. Neural implant recipients aren\'t hard to find — many are public about life-changing procedures.')

    # Step 2
    s2 = doc.add_paragraph()
    s2.add_run('Step 2: Initial Access. ').bold = True
    s2.add_run('A vulnerability in the wireless protocol — Bluetooth, WiFi, or the proprietary link to the external controller. Zero-days in wireless stacks are discovered regularly. The implant industry is small; security research resources are limited.')

    # Step 3
    s3 = doc.add_paragraph()
    s3.add_run('Step 3: Persistence. ').bold = True
    s3.add_run('The attacker establishes a foothold. They don\'t want to lose access if the device reboots or the patient moves out of range. They modify firmware, install a backdoor, or compromise the cloud service that manages device updates.')

    # Step 4
    s4 = doc.add_paragraph()
    s4.add_run('Step 4: Payload Deployment. ').bold = True
    s4.add_run('Now the ransomware activates. The device\'s therapeutic functions are disabled or degraded. A message is delivered — through the device\'s own interface if it has one, or through email/text linked to the patient\'s account.')

    # Step 5
    s5 = doc.add_paragraph()
    s5.add_run('Step 5: Extortion. ').bold = True
    s5.add_run('Pay or suffer. The attacker demands cryptocurrency. The victim faces a choice: lose the function the implant provides — possibly permanently — or pay the ransom with no guarantee of restoration.')

    add_separator(doc)

    # ========== SECTION: Why It's Worse ==========
    doc.add_heading('Why Neural Ransomware Is Worse Than Regular Ransomware', level=1)

    p3 = clean_text("""
    When ransomware hits your laptop, you lose access to files. It's painful, expensive, sometimes devastating — but it's recoverable. You can wipe the drive, restore from backup, buy a new machine.

    When ransomware hits your neural implant, you lose access to yourself.
    """)
    doc.add_paragraph(p3)

    # Point 1
    pt1 = doc.add_paragraph()
    pt1.add_run('You can\'t just "turn it off." ').bold = True
    pt1.add_run('Many implants manage critical functions. Deep brain stimulators control Parkinson\'s tremors. Cochlear implants provide hearing. Retinal implants provide vision. Turning them off isn\'t an inconvenience — it\'s a medical emergency.')

    # Point 2
    pt2 = doc.add_paragraph()
    pt2.add_run('You can\'t easily replace it. ').bold = True
    pt2.add_run('These devices require surgery to implant and surgery to remove. The brain may have adapted to the implant over months or years. Removal isn\'t just expensive — it may be medically risky.')

    # Point 3
    pt3 = doc.add_paragraph()
    pt3.add_run('The leverage is absolute. ').bold = True
    pt3.add_run('Regular ransomware holds your data hostage. Neural ransomware holds your body hostage. Your ability to see, hear, move, think. The psychological pressure is incomparable.')

    # Point 4
    pt4 = doc.add_paragraph()
    pt4.add_run('There\'s no "restore from backup." ').bold = True
    pt4.add_run('Your brain\'s neural pathways have been physically modified by the implant. Even if you could restore device firmware, you can\'t restore the biological changes that occurred during calibration.')

    add_separator(doc)

    # ========== SECTION: The Economics ==========
    doc.add_heading('The Economics Favor Attackers', level=1)

    p4 = clean_text("""
    Here's the uncomfortable math:

    A neural implant costs $50,000 to $150,000. Surgery adds another $50,000+. Insurance may or may not cover it. The patient has made an enormous investment — financial, physical, emotional.

    Now an attacker demands $10,000 in Bitcoin.

    What do you do? Fight it on principle while you can't see? Hire lawyers while your tremors return? Wait for the manufacturer to issue a patch while your quality of life collapses?

    Most people will pay. The attacker knows this.

    And unlike laptop ransomware, there's no IT department to call. No "restore from Time Machine." No buying a new one at Best Buy. The asymmetry is total.
    """)
    doc.add_paragraph(p4)

    add_separator(doc)

    # ========== SECTION: It Gets Worse ==========
    doc.add_heading('It Gets Worse: Beyond Simple Lockout', level=1)

    p5 = clean_text("""
    Locking the device is the simplest attack. But once an attacker has control of something inside your head, other possibilities emerge:
    """)
    doc.add_paragraph(p5)

    # Threat 1
    t1 = doc.add_paragraph()
    t1.add_run('Degradation attacks. ').bold = True
    t1.add_run('Instead of full lockout, the attacker slowly reduces functionality. Your vision gets slightly worse each day. Your tremor control becomes slightly less effective. You might not even realize it\'s an attack — until the ransom note arrives.')

    # Threat 2
    t2 = doc.add_paragraph()
    t2.add_run('Data exfiltration. ').bold = True
    t2.add_run('BCIs don\'t just write to the brain — they read from it. An attacker could record and sell your neural activity patterns. What are your thoughts worth to an advertiser? An employer? A government?')

    # Threat 3
    t3 = doc.add_paragraph()
    t3.add_run('Manipulation attacks. ').bold = True
    t3.add_run('If the implant can stimulate neural tissue, an attacker could theoretically influence mood, perception, or behavior. Pay us, or we\'ll make you feel afraid every time you try to leave your house.')

    # Threat 4
    t4 = doc.add_paragraph()
    t4.add_run('Destruction. ').bold = True
    t4.add_run('The nuclear option: deliberately damaging neural tissue through malicious stimulation patterns. "Pay or we permanently blind you." This crosses from extortion into something closer to assault — but the technical capability may exist.')

    add_separator(doc)

    # ========== SECTION: Current Defenses ==========
    doc.add_heading('Current Defenses Are Inadequate', level=1)

    p6 = clean_text("""
    I've spoken with people in the BCI industry. The security mindset is... developing.

    Most devices rely on "security through obscurity" — proprietary protocols that haven't been publicly analyzed. This is not security. It's delayed insecurity. Every proprietary protocol gets reverse-engineered eventually.

    Encryption exists but is often limited by power constraints. The implant runs on milliwatts. Strong cryptography burns energy. Trade-offs get made.

    Update mechanisms are particularly concerning. How do you patch firmware on something inside someone's skull? Over-the-air updates are convenient but expand the attack surface. Requiring hospital visits for every security patch is impractical.

    And there's no standard. No BCI equivalent of automotive cybersecurity regulations. No mandatory penetration testing. No bug bounty programs. The industry is moving fast and security is struggling to keep up.
    """)
    doc.add_paragraph(p6)

    add_separator(doc)

    # ========== SECTION: What Would Help ==========
    doc.add_heading('What Would Actually Help', level=1)

    p7 = clean_text("""
    This isn't hopeless. But it requires taking the threat seriously before it becomes a crisis.
    """)
    doc.add_paragraph(p7)

    # Defense 1
    d1 = doc.add_paragraph()
    d1.add_run('Hardware-enforced safety limits. ').bold = True
    d1.add_run('Analog circuits that physically prevent dangerous stimulation patterns, regardless of what firmware says. No software should be able to override hardware safety bounds. This is the last line of defense when everything else fails.')

    # Defense 2
    d2 = doc.add_paragraph()
    d2.add_run('Cryptographic device identity. ').bold = True
    d2.add_run('Each implant should have a unique, hardware-rooted cryptographic identity that cannot be cloned or spoofed. Commands must be signed by authorized keys. Unauthorized commands get rejected at the hardware level.')

    # Defense 3
    d3 = doc.add_paragraph()
    d3.add_run('Local-first architecture. ').bold = True
    d3.add_run('Critical functions should work without network connectivity. If the cloud service goes down — or gets compromised — the implant should continue functioning in a safe mode. No single point of failure outside the patient\'s body.')

    # Defense 4
    d4 = doc.add_paragraph()
    d4.add_run('Transparent security research. ').bold = True
    d4.add_run('The BCI industry needs to embrace security researchers, not threaten them with lawsuits. Bug bounty programs. Responsible disclosure policies. Published security audits. Obscurity isn\'t working.')

    # Defense 5
    d5 = doc.add_paragraph()
    d5.add_run('Regulatory requirements. ').bold = True
    d5.add_run('The FDA approves BCIs for safety and efficacy. Security should be part of that approval. Mandatory threat modeling. Required penetration testing. Incident reporting obligations. Make security a condition of market access.')

    add_separator(doc)

    # ========== SECTION: The Neural Firewall ==========
    doc.add_heading('This Is Why I\'m Building the Neural Firewall', level=1)

    p8 = clean_text("""
    Last week I introduced the Coherence Metric — a way to mathematically evaluate whether a neural signal should be trusted. But coherence is just one layer of defense.

    The ONI (Organic Neural Firewall) framework I'm developing treats the BCI as what it is: a network interface into your nervous system. And like any network interface, it needs a firewall.

    That firewall operates at the boundary between silicon and synapse. It inspects every signal in both directions. It validates cryptographic signatures. It checks coherence scores. It enforces rate limits and safety bounds. It logs everything for forensic analysis.

    Most importantly, it fails safe. If something seems wrong — if authentication fails, if coherence drops, if patterns look suspicious — the firewall blocks the signal and alerts the user. Better a false alarm than a successful attack.

    Neural ransomware is coming. The question isn't whether, but when. The time to build defenses is now — before the first victim's ransom note appears in their visual cortex.
    """)
    doc.add_paragraph(p8)

    add_separator(doc)

    # ========== CALL TO ACTION ==========
    doc.add_heading('What You Can Do', level=1)

    cta = clean_text("""
    If you're a patient or potential patient: Ask your device manufacturer about their security practices. What encryption do they use? How do they handle firmware updates? What happens if there's a security incident? You have a right to know.

    If you're in the BCI industry: Take this seriously. Hire security engineers. Commission penetration tests. Establish bug bounty programs. Don't wait for the first attack to make security a priority.

    If you're a policymaker: Start thinking about regulatory frameworks. BCIs are medical devices that are also networked computers. They need security requirements that reflect both identities.

    If you're a security researcher: This field needs you. The attack surface is enormous and largely unexplored. Your skills could literally protect people's minds.

    The brain is the last frontier of hacking. Let's make sure we're ready.
    """)
    doc.add_paragraph(cta)

    add_separator(doc)

    # ========== SERIES NOTE ==========
    series = doc.add_paragraph()
    series_run = series.add_run('This is the second article in a series on the ONI (Organic Neural Firewall) Framework. Previously: "Your Brain Has a Spam Filter. Can We Reverse-Engineer It?" Next week: "The Scale-Frequency Invariant — why f × S ≈ k holds from neurons to networks."')
    series_run.italic = True

    doc.add_paragraph()

    # Link to paper
    paper_link = doc.add_paragraph()
    paper_link.add_run('Read the technical deep-dive: ').bold = True
    paper_link.add_run('"Neural Ransomware: Attack Vectors and Defensive Architectures for Brain-Computer Interfaces" [link]')

    add_separator(doc)

    # ========== AUTHOR BIO ==========
    bio = doc.add_paragraph()
    bio_run = bio.add_run('Kevin L. Qi works at the intersection of cybersecurity, neuroscience, and AI governance. His background includes 15 years in cyber threat intelligence, biotech IT, and adversarial modeling. He\'s currently developing security frameworks for the bio-digital interfaces that don\'t exist yet — but will.')
    bio_run.italic = True

    add_separator(doc)

    # ========== TAGS ==========
    tags = doc.add_paragraph()
    tags.add_run('Tags: ').bold = True
    tags.add_run('#Cybersecurity #Ransomware #BrainComputerInterface #Neuralink #Privacy #Neuroscience #ONI')

    # Save
    output_path = '/Users/mac/Research/Articles/Week2_Neural_Ransomware_Medium.docx'
    doc.save(output_path)
    print(f'Medium blog post saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_medium_post()
