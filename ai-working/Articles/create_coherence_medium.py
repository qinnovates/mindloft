#!/usr/bin/env python3
"""
Create Medium Blog Post: The Coherence Metric
Formatted for Medium - No tables, accessible to general audience
Directs readers to detailed paper for technical depth
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_separator(doc):
    """Add a three-dot separator (Medium style)"""
    p = doc.add_paragraph()
    p.add_run('• • •')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_medium_post():
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========== TITLE ==========
    title = doc.add_heading('', 0)
    title_run = title.add_run('Your Brain Has a Spam Filter. Can We Reverse-Engineer It?')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Inside the math that could protect your mind from neural hackers.')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    add_separator(doc)

    # ========== HOOK ==========
    doc.add_paragraph("""Here's something that should keep you up at night:

When Neuralink sends a signal to your brain, your neurons can't tell the difference between that signal and one they generated themselves.

If the timing is right, the amplitude is right, and the frequency is right — your brain just... accepts it. No verification. No authentication. No "are you sure you want to allow this app to control your motor cortex?"

Evolution never anticipated signals arriving from silicon.""")

    doc.add_paragraph("""This is incredible when it helps a paralyzed patient move a cursor with their thoughts.

It's terrifying when you realize the same mechanism could let an attacker inject commands directly into your nervous system.""")

    add_separator(doc)

    # ========== THE INSIGHT ==========
    doc.add_heading('But Your Brain Isn\'t Completely Defenseless', level=1)

    doc.add_paragraph("""Here's what's fascinating: your brain actually does have a quality filter. It's just not designed for security — it's designed for signal integrity.

Every millisecond, your neurons are making decisions about which signals to trust and which to ignore. A spike that arrives at the wrong time gets filtered out. A signal that's too weak doesn't propagate. One that's too strong triggers protective mechanisms.

Your brain has been solving the "which signals are real?" problem for 500 million years. It just never had to solve it adversarially.""")

    doc.add_paragraph("""The question I've been obsessing over: can we formalize what the brain already knows? Can we turn its implicit quality filter into an explicit security check?""")

    add_separator(doc)

    # ========== THE METRIC ==========
    doc.add_heading('Introducing the Coherence Metric', level=1)

    doc.add_paragraph("""I've been developing a mathematical framework that attempts to quantify signal "trustworthiness" across three dimensions:""")

    p1 = doc.add_paragraph()
    p1.add_run('Timing. ').bold = True
    p1.add_run('Your neurons communicate through precisely timed oscillations — brain waves. Gamma rhythms pulse 30-100 times per second, and signals need to arrive at exactly the right phase to be processed. Miss the window by a few milliseconds and your brain\'s own gating mechanisms reject the signal. Random-phase attacks fail naturally.')

    p2 = doc.add_paragraph()
    p2.add_run('Pathway integrity. ').bold = True
    p2.add_run('Biological signals degrade as they travel through axons, across synapses, through dendrites. Each hop introduces a little noise, a little uncertainty. A signal that arrives too "clean" — with suspiciously low noise — might actually be a red flag. It bypassed the normal biological pathway.')

    p3 = doc.add_paragraph()
    p3.add_run('Amplitude. ').bold = True
    p3.add_run('Too weak and a signal won\'t trigger downstream neurons. Too strong and it can cause damage — excitotoxicity, receptor saturation, tissue harm. Your brain maintains homeostatic balance through mechanisms we\'re only beginning to understand. Signals outside the expected amplitude range get flagged.')

    add_separator(doc)

    # ========== THE FORMULA ==========
    doc.add_heading('One Equation, Three Dimensions', level=1)

    doc.add_paragraph('Combine these three factors and you get what I call the Coherence Metric:')

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = formula.add_run('Cₛ = e^(−(σ²ᵩ + σ²τ + σ²ᵧ))')
    f_run.font.size = Pt(16)
    f_run.bold = True

    doc.add_paragraph("""Don't worry about the Greek letters. Here's what matters:

When a signal has perfect timing, perfect pathway integrity, and perfect amplitude — coherence equals 1. Full trust.

As any of those factors degrade — timing jitter, transmission noise, amplitude fluctuation — coherence drops toward zero. Increasing suspicion.

The exponential decay isn't arbitrary. Neural systems exhibit threshold behaviors. A signal doesn't gradually become "less trusted" — it either crosses the threshold for propagation or it doesn't. The math captures that biological reality.""")

    add_separator(doc)

    # ========== WHY IT MATTERS ==========
    doc.add_heading('Why This Matters for Brain-Computer Interfaces', level=1)

    doc.add_paragraph("""Imagine a Neural Firewall — a security layer that sits between the digital world and your neural tissue, inspecting every signal in both directions.

For incoming commands, it asks: Does this signal's coherence score fall within biological norms? Is the timing aligned with ongoing brain rhythms? Is the amplitude within safe bounds?

If any check fails: reject, log, alert.

For outgoing signals, it asks: Is this normal neural activity or something anomalous? Should we strip sensitive information before transmission? Is everything encrypted before it hits Bluetooth?""")

    doc.add_paragraph("""The hard part: this firewall has to run on a chip that draws less power than a hearing aid. Neuralink's implant runs on 25 milliwatts — that's nothing. The security layer gets maybe 3-5 milliwatts to work with.

But here's the thing — the coherence calculation is surprisingly efficient. You're not analyzing raw signals; you're tracking statistics. Means and variances. A few hundred microseconds of latency. It's doable.""")

    add_separator(doc)

    # ========== THE CATCH ==========
    doc.add_heading('The Catch (Because There\'s Always a Catch)', level=1)

    doc.add_paragraph("""This framework isn't bulletproof. A sophisticated attacker with read access to your brain's electrical activity could potentially synchronize their malicious signals to your ongoing rhythms — achieving high coherence while still being harmful.

The metric detects abnormal signals. It doesn't guarantee detection of all malicious ones.

It's also not empirically validated. The math is grounded in neuroscience literature, but we need experiments — animal studies, clinical correlations — to prove these metrics actually predict what the brain accepts versus rejects.""")

    doc.add_paragraph("""But here's why I'm publishing anyway: we need shared vocabulary.

Before we can defend bio-digital interfaces, we need to be able to talk about what we're defending. What's the coherence threshold? Which variance component is most vulnerable? How do we detect phase-synchronized attacks?

These are the questions we should be asking now — before brain-computer interfaces are in millions of heads.""")

    add_separator(doc)

    # ========== BIGGER PICTURE ==========
    doc.add_heading('Part of Something Larger', level=1)

    doc.add_paragraph("""The coherence metric is one piece of a larger framework I've been developing called ONI — the Organic Neural Firewall.

Think of it as the OSI model extended into biology. Seven traditional network layers, plus seven more that characterize what happens when signals cross from silicon into neural tissue — from ion channels to oscillations to working memory to identity itself.

Each layer has its own attack surfaces. Each requires its own defenses. The coherence metric operates at Layers 8-10 — the neural interface domain where digital meets biological.

The goal isn't to solve everything. It's to create a scaffold that neuroscientists, security engineers, and ethicists can stress-test, criticize, and improve.""")

    add_separator(doc)

    # ========== CALL TO ACTION ==========
    doc.add_heading('What\'s Next', level=1)

    doc.add_paragraph("""I've published a detailed technical paper expanding everything here — formal mathematical derivations, security analysis, hardware implementation proposals, and honest discussion of limitations.

If you're a neuroscientist: tell me what's wrong with the biological assumptions.

If you're a security engineer: tell me how you'd break this.

If you're building BCIs: tell me what constraints I'm missing.

The brain's firewall is not optional. It's the minimum viable security for any system that touches living neural tissue. Let's build it right.""")

    add_separator(doc)

    # ========== SERIES NOTE ==========
    series = doc.add_paragraph()
    series_run = series.add_run('This is the first article in a series on the ONI (Organic Neural Firewall) Framework. Next week: "Neural Ransomware Isn\'t Science Fiction" — a technical breakdown of how attackers could hold your implant hostage.')
    series_run.italic = True

    doc.add_paragraph()

    # Link to paper
    paper_link = doc.add_paragraph()
    paper_link.add_run('Read the full technical paper: ').bold = True
    paper_link.add_run('"The Coherence Metric for Neural Signal Integrity" [link]')

    add_separator(doc)

    # ========== AUTHOR BIO ==========
    bio = doc.add_paragraph()
    bio_run = bio.add_run('Kevin L. Qi works at the intersection of cybersecurity, neuroscience, and AI governance. His background includes 15 years in cyber threat intelligence, biotech IT, and adversarial modeling. He\'s currently developing security frameworks for the bio-digital interfaces that don\'t exist yet — but will.')
    bio_run.italic = True

    add_separator(doc)

    # ========== TAGS ==========
    tags = doc.add_paragraph()
    tags.add_run('Tags: ').bold = True
    tags.add_run('#Neuroscience #BrainComputerInterface #Cybersecurity #Neuralink #Privacy #AI #ONI')

    # Save
    output_path = '/Users/mac/Research/Articles/Week1_Coherence_Metric_Medium.docx'
    doc.save(output_path)
    print(f'Medium blog post saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_medium_post()
