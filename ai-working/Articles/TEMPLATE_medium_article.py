#!/usr/bin/env python3
"""
TEMPLATE: Medium Blog Post Generator for ONI Framework Series
=============================================================

DESIGN PRINCIPLES:
1. Catchy title + subtitle that hooks readers
2. No tables (Medium doesn't render them well)
3. Accessible to non-experts - explain concepts simply
4. Complete story: Problem → Insight → Solution → Limitations → Call to Action
5. Entice readers to explore the technical paper and other articles
6. Use bold text for key terms, not complex formatting

Usage:
1. Copy this file and rename for your article
2. Update the ARTICLE_CONFIG dictionary with your content
3. Run the script to generate .docx
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import textwrap

# ============================================================
# ARTICLE CONFIGURATION - UPDATE THIS SECTION FOR EACH ARTICLE
# ============================================================

ARTICLE_CONFIG = {
    # Basic metadata
    'title': 'Your Catchy Title Here',  # Make it intriguing, not academic
    'subtitle': 'A compelling hook that promises value to the reader.',
    'author': 'Kevin L. Qi',
    'output_filename': 'Article_Name_Medium.docx',

    # Tags for Medium (5-7 relevant tags)
    'tags': ['#Neuroscience', '#BrainComputerInterface', '#Cybersecurity', '#Neuralink', '#AI', '#ONI'],

    # Author bio (keep consistent across series)
    'author_bio': "Kevin L. Qi works at the intersection of cybersecurity, neuroscience, and AI governance. His background includes 15 years in cyber threat intelligence, biotech IT, and adversarial modeling. He's currently developing security frameworks for the bio-digital interfaces that don't exist yet — but will.",

    # Series info
    'series_note': 'This article is part of a series on the ONI (Organic Neural Firewall) Framework.',
    'next_article_teaser': 'Next week: "Title of Next Article" — brief enticing description.',

    # Link to detailed paper
    'paper_title': 'Full Paper Title Here',
    'paper_link': '[link]',

    # Acknowledgements (included in papers, optional for Medium posts)
    'include_acknowledgements': True,  # Set to False for Medium posts if too long
    'acknowledgements': """The author wishes to acknowledge the support and inspiration of colleagues and mentors in the development of this work. Writing and structural assistance was provided by Claude Code (Anthropic, 2025). Diagramming, visualization, and threat-model representation assistance was provided by ChatGPT (OpenAI, 2025). All ideas, analyses, interpretations, and conclusions presented herein are entirely the author's own.""",

    # AI Tool Citations (APA 7th Edition format)
    'ai_citations': [
        'Anthropic. (2025). Claude Code (Version 4.5) [Large language model]. https://www.anthropic.com/claude',
        'OpenAI. (2025). ChatGPT (Version 4) [Large language model]. https://chat.openai.com',
    ],

    # ========== CONTENT SECTIONS ==========
    # Structure: Hook → Problem → Insight → Solution → Catch → Bigger Picture → CTA

    'hook': """Your opening hook. Make it visceral, surprising, or thought-provoking.

Start with something that creates tension or curiosity. Use current events, striking facts, or relatable scenarios.

End with a question or implication that the article will address.""",

    'sections': [
        {
            'heading': 'The Problem (Make It Real)',
            'paragraphs': [
                "Explain the problem in accessible terms. Why should readers care?",
                "Use concrete examples. Avoid jargon or explain it immediately.",
                "Build tension — what's at stake?",
            ]
        },
        {
            'heading': 'The Insight (Your Unique Angle)',
            'paragraphs': [
                "What did you discover or realize that others haven't?",
                "Frame it as a story of discovery if possible.",
                "This is where you differentiate from other articles on the topic.",
            ]
        },
        {
            'heading': 'The Solution (Keep It Simple)',
            'paragraphs': [
                "Explain your framework/approach in plain language.",
                "Use bold text for key concepts: **Timing.** Explanation follows...",
                "If there's a formula, show it but immediately explain what it means intuitively.",
            ],
            # Optional formula - will be centered and bolded
            'formula': 'Your formula here (optional)',
            'formula_explanation': "Don't worry about the math. Here's what matters: [plain language explanation]"
        },
        {
            'heading': 'Why This Matters',
            'paragraphs': [
                "Connect back to the real-world implications.",
                "Paint a picture of the future this enables or prevents.",
                "Make it tangible — what changes if this works?",
            ]
        },
        {
            'heading': 'The Catch (Be Honest)',
            'paragraphs': [
                "What are the limitations? What doesn't work yet?",
                "Being honest about gaps builds credibility.",
                "Frame limitations as opportunities for collaboration.",
            ]
        },
        {
            'heading': 'Part of Something Larger',
            'paragraphs': [
                "Connect this article to the broader ONI framework.",
                "Tease other aspects readers might want to explore.",
                "Position this as one piece of a larger puzzle.",
            ]
        },
    ],

    'call_to_action': {
        'heading': "What's Next",
        'content': """Specific asks for different audiences:

If you're a [role]: tell me [specific feedback you want].

If you're a [role]: tell me [specific feedback you want].

End with a memorable closing line that reinforces the core message."""
    }
}

# ============================================================
# DOCUMENT GENERATOR - DO NOT MODIFY BELOW THIS LINE
# ============================================================

def clean_text(text):
    """Remove leading whitespace from each line while preserving paragraph breaks"""
    return textwrap.dedent(text).strip()

def add_separator(doc):
    """Add a three-dot separator (Medium style)"""
    p = doc.add_paragraph()
    p.add_run('• • •')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_medium_post(config=ARTICLE_CONFIG):
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # TITLE
    title = doc.add_heading('', 0)
    title_run = title.add_run(config['title'])
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(config['subtitle'])
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    add_separator(doc)

    # HOOK
    doc.add_paragraph(clean_text(config['hook']))

    add_separator(doc)

    # MAIN SECTIONS
    for section in config['sections']:
        doc.add_heading(section['heading'], level=1)

        for para in section.get('paragraphs', []):
            cleaned_para = clean_text(para)
            # Handle bold text marked with **text**
            if '**' in cleaned_para:
                p = doc.add_paragraph()
                parts = cleaned_para.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
            else:
                doc.add_paragraph(cleaned_para)

        # Optional formula
        if 'formula' in section and section['formula']:
            doc.add_paragraph()
            formula = doc.add_paragraph()
            formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
            f_run = formula.add_run(section['formula'])
            f_run.font.size = Pt(16)
            f_run.bold = True

            if 'formula_explanation' in section:
                doc.add_paragraph(clean_text(section['formula_explanation']))

        add_separator(doc)

    # CALL TO ACTION
    cta = config['call_to_action']
    doc.add_heading(cta['heading'], level=1)
    doc.add_paragraph(clean_text(cta['content']))

    add_separator(doc)

    # SERIES NOTE
    series = doc.add_paragraph()
    series_run = series.add_run(config['series_note'] + ' ' + config['next_article_teaser'])
    series_run.italic = True

    doc.add_paragraph()

    # Link to paper
    paper_link = doc.add_paragraph()
    paper_link.add_run('Read the full technical paper: ').bold = True
    paper_link.add_run(f'"{config["paper_title"]}" {config["paper_link"]}')

    add_separator(doc)

    # AUTHOR BIO
    bio = doc.add_paragraph()
    bio_run = bio.add_run(config['author_bio'])
    bio_run.italic = True

    add_separator(doc)

    # ACKNOWLEDGEMENTS (if enabled)
    if config.get('include_acknowledgements', False) and config.get('acknowledgements'):
        doc.add_heading('Acknowledgements', level=1)
        doc.add_paragraph(clean_text(config['acknowledgements']))
        add_separator(doc)

    # AI CITATIONS (if provided)
    if config.get('ai_citations'):
        doc.add_heading('AI Tools Citation (APA 7th Edition)', level=1)
        for citation in config['ai_citations']:
            doc.add_paragraph(clean_text(citation))
        add_separator(doc)

    # TAGS
    tags = doc.add_paragraph()
    tags.add_run('Tags: ').bold = True
    tags.add_run(' '.join(config['tags']))

    # Save
    output_path = f"/Users/mac/Research/Articles/{config['output_filename']}"
    doc.save(output_path)
    print(f"Article saved to: {output_path}")
    return output_path

# ============================================================
# EXAMPLE: How to use bold text in paragraphs
# ============================================================
#
# In your paragraphs, wrap text in **double asterisks** to make it bold:
#
# "**Timing.** Your neurons communicate through precisely timed oscillations."
#
# This will render as:
# Timing. Your neurons communicate through precisely timed oscillations.
# (with "Timing." in bold)
#
# ============================================================
# ACKNOWLEDGEMENTS AND AI CITATIONS
# ============================================================
#
# The template includes:
# - Acknowledgements section crediting AI assistance
# - APA 7th Edition citations for Claude and ChatGPT
#
# For Medium posts: Set 'include_acknowledgements': False if too long
# For research papers: Keep 'include_acknowledgements': True
#
# APA 7th Edition format for AI tools:
# Author. (Year). Tool Name (Version) [Description]. URL
#
# Example:
# Anthropic. (2025). Claude Code (Version 4.5) [Large language model]. https://www.anthropic.com/claude
#
# ============================================================

if __name__ == '__main__':
    create_medium_post()
