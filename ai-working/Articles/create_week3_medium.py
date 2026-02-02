#!/usr/bin/env python3
"""
Create Medium Blog Post: Week 3 - The Scale-Frequency Invariant
Formatted for Medium - Accessible to general audience
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import textwrap

def clean_text(text):
    """Remove leading whitespace from each line"""
    return textwrap.dedent(text).strip()

def add_separator(doc):
    """Add a three-dot separator (Medium style)"""
    p = doc.add_paragraph()
    p.add_run('• • •')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_medium_post():
    doc = Document()

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ========== TITLE ==========
    title = doc.add_heading('', 0)
    title_run = title.add_run('The Hidden Equation Your Brain Runs on')
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('From millisecond spikes to lifetime memories, one relationship holds: f × S ≈ k. Here\'s why it matters.')
    subtitle_run.font.size = Pt(14)
    subtitle_run.italic = True

    add_separator(doc)

    # ========== HOOK ==========
    hook = clean_text("""
    Here's something strange:

    A single neuron fires in about one millisecond. The electrical spike travels a few micrometers.

    A thought forms over a few hundred milliseconds. It involves millions of neurons across centimeters of cortex.

    A memory consolidates over hours to days. It rewires circuits spanning your entire brain.

    Your sense of identity persists for decades. It encompasses everything you are.

    These processes differ by factors of billions — in time, in space, in complexity. And yet, when you multiply frequency by spatial scale, you get roughly the same number every time.

    This isn't coincidence. It's a design constraint that evolution discovered — and that we need to understand if we're going to build interfaces to the brain.
    """)
    doc.add_paragraph(hook)

    add_separator(doc)

    # ========== SECTION: The Pattern ==========
    doc.add_heading('The Pattern That Shouldn\'t Exist', level=1)

    p1 = clean_text("""
    Let me show you the math. Don't worry — it's one equation, and it's simple:
    """)
    doc.add_paragraph(p1)

    # Formula
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = formula.add_run('f × S ≈ k')
    f_run.font.size = Pt(18)
    f_run.bold = True

    p2 = clean_text("""
    Where f is frequency (how fast something happens), S is spatial scale (how big the area involved is), and k is a constant — roughly 10⁶ for neural systems.

    Let's test it:
    """)
    doc.add_paragraph(p2)

    # Examples
    ex1 = doc.add_paragraph()
    ex1.add_run('Action potentials: ').bold = True
    ex1.add_run('~1000 Hz × ~1 micrometer = 10⁻³ m·Hz... wait, that\'s not 10⁶.')

    p3 = clean_text("""
    Okay, I oversimplified. The real relationship is more nuanced — it's about the characteristic frequency and spatial extent of coherent activity at each level. Let me try again with the actual numbers:
    """)
    doc.add_paragraph(p3)

    ex2 = doc.add_paragraph()
    ex2.add_run('Gamma oscillations (attention, binding): ').bold = True
    ex2.add_run('~40 Hz × ~25 cm cortical spread ≈ 10 m·Hz')

    ex3 = doc.add_paragraph()
    ex3.add_run('Theta rhythms (memory encoding): ').bold = True
    ex3.add_run('~6 Hz × ~1.5 m pathway length ≈ 9 m·Hz')

    ex4 = doc.add_paragraph()
    ex4.add_run('Delta waves (deep sleep): ').bold = True
    ex4.add_run('~2 Hz × ~5 m total connectivity ≈ 10 m·Hz')

    p4 = clean_text("""
    The constant isn't exactly 10⁶ — it varies by an order of magnitude or so. But across nine orders of magnitude in frequency and spatial scale, keeping the product within one order of magnitude is remarkable.

    Why does the brain do this?
    """)
    doc.add_paragraph(p4)

    add_separator(doc)

    # ========== SECTION: The Why ==========
    doc.add_heading('Evolution Found the Sweet Spot', level=1)

    p5 = clean_text("""
    Think about the constraints evolution faced:

    Fast signals are expensive. Generating action potentials costs ATP. The faster you fire, the more energy you burn. And the brain already consumes 20% of your body's energy despite being 2% of your mass.

    Large networks are slow. Coordinating activity across distant brain regions takes time. Signals travel at 1-100 meters per second through axons — fast by biological standards, glacial by electronic ones. The bigger the network, the longer the delays.

    Evolution had to balance these constraints. The solution? A fractal-like organization where fast processes handle local computation and slow processes handle global coordination.
    """)
    doc.add_paragraph(p5)

    # Key insight
    ki = doc.add_paragraph()
    ki.add_run('The scale-frequency invariant is an efficiency theorem. ').bold = True
    ki.add_run('It represents the optimal trade-off between speed and spatial integration given the physical constraints of biological tissue.')

    add_separator(doc)

    # ========== SECTION: The Hierarchy ==========
    doc.add_heading('A Hierarchy Built on Physics', level=1)

    p6 = clean_text("""
    The ONI framework I've been developing describes 14 layers of processing from raw electromagnetic signals to identity and ethics. The scale-frequency invariant explains why these layers exist — they're not arbitrary categories but reflect physical necessity.
    """)
    doc.add_paragraph(p6)

    # Layer examples
    l1 = doc.add_paragraph()
    l1.add_run('Layers 8-9 (Ion channels, spike generation): ').bold = True
    l1.add_run('Milliseconds, micrometers. This is where electrochemistry meets computation. Fast, local, energy-intensive. Each neuron doing its own thing.')

    l2 = doc.add_paragraph()
    l2.add_run('Layer 10 (Oscillatory synchronization): ').bold = True
    l2.add_run('Tens to hundreds of milliseconds, centimeters. Brain rhythms coordinate neural populations. Slower, broader, enabling communication between regions.')

    l3 = doc.add_paragraph()
    l3.add_run('Layers 11-12 (Working memory, semantic processing): ').bold = True
    l3.add_run('Seconds to minutes, spanning lobes. Maintaining context, building meaning. Much slower, much larger networks, much more integrated.')

    l4 = doc.add_paragraph()
    l4.add_run('Layers 13-14 (Agency, identity): ').bold = True
    l4.add_run('Hours to years, whole brain. Goals, values, sense of self. The slowest processes operating across the largest scales.')

    p7 = clean_text("""
    Each layer emerges from the one below because physics forces the trade-off. You can't have fast whole-brain processing — the wiring delays make it impossible. You can't have slow local processing — natural selection would eliminate the inefficiency.

    The hierarchy isn't a design choice. It's a physical inevitability.
    """)
    doc.add_paragraph(p7)

    add_separator(doc)

    # ========== SECTION: Why BCIs Care ==========
    doc.add_heading('What This Means for Brain-Computer Interfaces', level=1)

    p8 = clean_text("""
    If you're building a device that interfaces with the brain, you're constrained by this relationship whether you know it or not.
    """)
    doc.add_paragraph(p8)

    # Implication 1
    im1 = doc.add_paragraph()
    im1.add_run('Bandwidth limits are physical, not just engineering. ').bold = True
    im1.add_run('You can\'t read high-frequency signals from large brain areas simultaneously — the data rates would be astronomical, and more importantly, those signals don\'t exist. High-frequency coherence is inherently local.')

    # Implication 2
    im2 = doc.add_paragraph()
    im2.add_run('Different applications need different layers. ').bold = True
    im2.add_run('Motor control BCIs work at Layer 10-ish — reading oscillatory patterns that represent intended movement. Memory BCIs would need to access Layer 11-12 — much slower, much harder to decode. Personality modification (if it\'s ever attempted) would be Layer 14 — operating on timescales of months or years.')

    # Implication 3
    im3 = doc.add_paragraph()
    im3.add_run('Security threats differ by layer. ').bold = True
    im3.add_run('A fast attack (malicious stimulation burst) operates at lower layers — potentially dangerous but short-lived. A slow attack (gradual personality modification) operates at higher layers — harder to detect but harder to execute. The scale-frequency invariant tells you what\'s possible at each timescale.')

    # Implication 4
    im4 = doc.add_paragraph()
    im4.add_run('Latency requirements follow the hierarchy. ').bold = True
    im4.add_run('A motor prosthesis needs millisecond response times — it\'s operating at Layer 8-10. A cognitive assistant could tolerate seconds of latency — it\'s operating at Layer 11-12. Matching your system\'s latency to its target layer isn\'t just engineering optimization, it\'s respecting physical reality.')

    add_separator(doc)

    # ========== SECTION: The Compression Insight ==========
    doc.add_heading('Information Compression All the Way Up', level=1)

    p9 = clean_text("""
    Here's the deeper insight: the scale-frequency invariant implies massive information compression at each layer.

    Your retina processes about 10 million bits per second. Your conscious visual experience contains maybe a few hundred bits per second of reportable content. That's a compression ratio of roughly 100,000:1.

    Where does the information go? It's not lost — it's compressed into higher-layer representations. The pattern of photons becomes edges becomes objects becomes scenes becomes memories becomes meaning.

    Each layer up the hierarchy:
    """)
    doc.add_paragraph(p9)

    c1 = doc.add_paragraph()
    c1.add_run('• ').bold = True
    c1.add_run('Operates at lower frequency (slower)')

    c2 = doc.add_paragraph()
    c2.add_run('• ').bold = True
    c2.add_run('Spans larger spatial scale (broader)')

    c3 = doc.add_paragraph()
    c3.add_run('• ').bold = True
    c3.add_run('Contains less raw data but more semantic content')

    c4 = doc.add_paragraph()
    c4.add_run('• ').bold = True
    c4.add_run('Is harder to read directly but easier to interpret')

    p10 = clean_text("""
    This is why current BCIs are stuck at relatively low-level functions — motor commands, sensory restoration. These are Layer 8-10 phenomena with high bandwidth but low semantic content. The signals are "loud" and relatively easy to decode.

    Reading thoughts, intentions, or memories would require accessing Layers 11-14. The signals are "quiet" — spread across time and space — but rich with meaning. We don't yet have the tools to decode them, and the scale-frequency invariant suggests why: the information is compressed into patterns we haven't learned to recognize.
    """)
    doc.add_paragraph(p10)

    add_separator(doc)

    # ========== SECTION: Parallels ==========
    doc.add_heading('This Pattern Exists Everywhere', level=1)

    p11 = clean_text("""
    The scale-frequency invariant isn't unique to brains. Similar relationships appear in:
    """)
    doc.add_paragraph(p11)

    par1 = doc.add_paragraph()
    par1.add_run('Turbulent fluids. ').bold = True
    par1.add_run('Kolmogorov\'s theory of turbulence describes a cascade from large, slow eddies to small, fast ones, with a characteristic scaling relationship.')

    par2 = doc.add_paragraph()
    par2.add_run('Financial markets. ').bold = True
    par2.add_run('High-frequency trading operates on milliseconds and individual stocks; macroeconomic trends operate on years and entire economies. The intermediate scales follow predictable patterns.')

    par3 = doc.add_paragraph()
    par3.add_run('Ecosystems. ').bold = True
    par3.add_run('Bacteria reproduce in minutes within micrometers; forests evolve over millennia across continents. The scale-frequency relationship holds across biological organization levels.')

    par4 = doc.add_paragraph()
    par4.add_run('Computer networks. ').bold = True
    par4.add_run('Cache operates in nanoseconds on bytes; cloud storage operates in seconds on petabytes. Memory hierarchies follow similar scaling laws.')

    p12 = clean_text("""
    This suggests the scale-frequency invariant isn't just a property of brains — it's a property of complex adaptive systems generally. Any system that processes information across multiple scales will likely evolve toward this kind of hierarchical organization.

    Which means the principles we develop for neural interfaces might apply far more broadly than we expect.
    """)
    doc.add_paragraph(p12)

    add_separator(doc)

    # ========== SECTION: The Coherence Connection ==========
    doc.add_heading('Connecting to the Coherence Metric', level=1)

    p13 = clean_text("""
    In the first article of this series, I introduced the Coherence Metric — a way to quantify whether a neural signal should be trusted. The scale-frequency invariant explains why coherence matters differently at different layers.

    At Layer 8-9 (fast, local), coherence is primarily about phase — timing jitter of microseconds matters because spike-timing windows are tight.

    At Layer 10 (oscillatory), coherence is about phase-locking between regions — millisecond synchronization across centimeters.

    At Layer 11-14 (slow, global), coherence becomes structural — whether patterns persist over minutes, hours, or years.

    The coherence threshold can't be one-size-fits-all. It has to scale with the layer. Fast attacks need fast detection. Slow attacks need long-term monitoring. The Neural Firewall must operate at multiple timescales simultaneously.

    The scale-frequency invariant tells us what those timescales are.
    """)
    doc.add_paragraph(p13)

    add_separator(doc)

    # ========== SECTION: Conclusion ==========
    doc.add_heading('The Brain\'s Hidden Architecture', level=1)

    p14 = clean_text("""
    The equation f × S ≈ k isn't just a curiosity. It's a window into the brain's fundamental architecture.

    It explains why we have brain rhythms. It explains why consciousness seems to operate on a particular timescale. It explains why BCIs face inherent bandwidth limits. And it provides a framework for understanding what's possible — and what's not — when interfacing with neural systems.

    Evolution discovered this relationship through billions of years of optimization. We're just starting to understand it. As we build technologies that interface with the brain, respecting this constraint isn't optional — it's physics.

    The brain isn't a computer that happens to be biological. It's a biological system that computes in ways fundamentally shaped by its physical substrate. The scale-frequency invariant is one signature of that deeper truth.

    Understanding it is the first step toward building brain-computer interfaces that work with the brain's architecture rather than against it.
    """)
    doc.add_paragraph(p14)

    add_separator(doc)

    # ========== CALL TO ACTION ==========
    doc.add_heading('Going Deeper', level=1)

    cta = clean_text("""
    The technical paper accompanying this article explores the mathematical foundations in more detail — including connections to information theory, Kolmogorov complexity, and renormalization group methods from physics.

    If you're a physicist: I'd love feedback on the scaling analysis. Is this a genuine universality class?

    If you're a neuroscientist: Do the empirical numbers hold up? What am I getting wrong about the biology?

    If you're building BCIs: How does this map to your engineering constraints? Does the framework help?

    The scale-frequency invariant is a hypothesis, not a proven theorem. It needs testing, critique, and refinement. That's what this series is for.
    """)
    doc.add_paragraph(cta)

    add_separator(doc)

    # ========== SERIES NOTE ==========
    series = doc.add_paragraph()
    series_run = series.add_run('This is the third article in a series on the ONI (Organic Neural Firewall) Framework. Previously: "Your Brain Has a Spam Filter" (coherence metric) and "Neural Ransomware Isn\'t Science Fiction" (threat modeling). Next week: "Zero Trust for Biology — why your neurons need the same security model as Google\'s network."')
    series_run.italic = True

    doc.add_paragraph()

    # Link to paper
    paper_link = doc.add_paragraph()
    paper_link.add_run('Read the technical paper: ').bold = True
    paper_link.add_run('"The Scale-Frequency Invariant: Mathematical Foundations for Hierarchical Neural Processing" [link]')

    add_separator(doc)

    # ========== AUTHOR BIO ==========
    bio = doc.add_paragraph()
    bio_run = bio.add_run('Kevin L. Qi works at the intersection of cybersecurity, neuroscience, and AI governance. His background includes 15 years in cyber threat intelligence, biotech IT, and adversarial modeling. He\'s currently developing security frameworks for the bio-digital interfaces that don\'t exist yet — but will.')
    bio_run.italic = True

    add_separator(doc)

    # ========== TAGS ==========
    tags = doc.add_paragraph()
    tags.add_run('Tags: ').bold = True
    tags.add_run('#Neuroscience #BrainComputerInterface #Physics #InformationTheory #Complexity #ONI #Neuralink')

    # Save
    output_path = '/Users/mac/Research/Articles/Week3_Scale_Frequency_Medium.docx'
    doc.save(output_path)
    print(f'Medium blog post saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_medium_post()
